# app.py
# ---------------------------------------------------------
# Food Blog Article Generator with Auto Internal Linking
# + Long-form multi-request mode for robust SEO articles
# + Publish to WordPress (REST API with Application Passwords)
# + Auto-create & embed Tasty Recipe (CPT) via shortcode block
#   (fallback to clean HTML card when CPT isn't available)
# ---------------------------------------------------------

import os
import io
import re
import json
import html
import datetime
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any

import streamlit as st
import streamlit.components.v1 as components
from openai import OpenAI

# Networking for WordPress
import requests
from requests.auth import HTTPBasicAuth

# DOCX export with real hyperlinks
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

# XML sitemap parsing
import xml.etree.ElementTree as ET

# Optional Markdown->HTML conversion
try:
    import markdown as mdlib  # pip install markdown
except Exception:
    mdlib = None


def md_to_html(text: str) -> str:
    """Convert Markdown text to HTML using python-markdown if available, otherwise a basic regex fallback."""
    if mdlib:
        try:
            return mdlib.markdown(text)
        except Exception:
            pass
    # Fallback: basic markdown to HTML conversion
    html = text
    # Convert headings
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    # Convert bold text
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    # Convert double line breaks to paragraphs
    html = re.sub(r'\n\n+', '</p><p>', html)
    html = f'<p>{html}</p>'
    html = html.replace('<p></p>', '')
    return html

# ---------------------------------------------------------
# Settings persistence
# ---------------------------------------------------------
SETTINGS_FILE = Path(__file__).with_name("settings.json")

def load_settings_into_session():
    try:
        # Load environment variables from .env if available
        try:
            from dotenv import load_dotenv
            load_dotenv()
        except Exception:
            # Fallback: manually load .env if python-dotenv is not installed
            try:
                env_path = Path(__file__).with_name(".env")
                if env_path.exists():
                    for line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
                        s = line.strip()
                        if not s or s.startswith("#"):
                            continue
                        if s.startswith("export "):
                            s = s[len("export "):].strip()
                        if "=" not in s:
                            continue
                        k, v = s.split("=", 1)
                        k = k.strip()
                        v = v.strip()
                        # strip optional quotes
                        if (v.startswith('"') and v.endswith('"')) or (v.startswith("'") and v.endswith("'")):
                            v = v[1:-1]
                        os.environ.setdefault(k, v)
            except Exception:
                pass
        if st.session_state.get("_settings_loaded"):
            return
        if SETTINGS_FILE.exists():
            data = json.loads(SETTINGS_FILE.read_text(encoding="utf-8", errors="ignore"))
            for k, v in data.items():
                if k == "API_KEY":
                    continue
                st.session_state[k] = v
        st.session_state["_settings_loaded"] = True
    except Exception as e:
        st.warning(f"Failed to load settings: {e}")

def save_current_settings():
    try:
        keys = [
            # auth (do NOT persist API keys)
            # "API_KEY",  # removed: use .env only

            # social + CTA
            "fb_url","pin_url","append_cta",

            # linking / sitemap
            "link_style","default_path","sitemap_url","sitemap_urls",
            "max_links","per_paragraph_max","link_headings","skip_lines_with_links",
            "hint_model_with_phrases",

            # generation
            "use_multi_call","target_words","model_name","temperature",

            # wordpress
            "wp_site","wp_user","wp_app_pw","wp_status","wp_categories",
            "wp_excerpt","feat_url","auto_post",

            # identity
            "writer_identity_name",
        ]
        data = {k: st.session_state.get(k) for k in keys}
        SETTINGS_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8", errors="ignore")
        st.success("Settings saved.")
    except Exception as e:
        st.error(f"Failed to save settings: {e}")

# Load before UI
load_settings_into_session()


# ---------------------------------------------------------
# OpenAI client
# ---------------------------------------------------------
def get_client():
    key = os.getenv("API_KEY")
    if not key:
        # Lazy-load .env now in case settings haven't been loaded yet
        try:
            from dotenv import load_dotenv
            load_dotenv()
        except Exception:
            # Manual .env parsing fallback
            try:
                env_path = Path(__file__).with_name(".env")
                if env_path.exists():
                    for line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
                        s = line.strip()
                        if not s or s.startswith("#"):
                            continue
                        if s.startswith("export "):
                            s = s[len("export ") :].strip()
                        if "=" not in s:
                            continue
                        k, v = s.split("=", 1)
                        k = k.strip()
                        v = v.strip()
                        if (v.startswith('"') and v.endswith('"')) or (v.startswith("'") and v.endswith("'")):
                            v = v[1:-1]
                        os.environ.setdefault(k, v)
            except Exception:
                pass
        # Try again after loading
        key = os.getenv("API_KEY")
    if not key:
        try:
            st.sidebar.error("API_KEY not found. Add it to your .env file.")
        except Exception:
            pass
        raise RuntimeError("API_KEY missing")
    return OpenAI(api_key=key)

# ---------------------------------------------------------
# Utilities
# ---------------------------------------------------------
def _normalize_url(u: str) -> str:
    u = (u or "").strip()
    if not u:
        return ""
    if not (u.startswith("http://") or u.startswith("https://")):
        u = "https://" + u
    return u

def slug_to_title(slug: str) -> str:
    s = slug.strip("/").split("/")[-1]
    s = re.sub(r"[-_]+", " ", s).strip()
    s = re.sub(r"\s+", " ", s)
    return s.title()

def to_slug(text: str) -> str:
    s = re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE)
    s = re.sub(r"\s+", "-", s).strip("-").lower()
    return s[:90] or "post"

def excerpt_from_text(html_or_md: str, max_words: int = 40) -> str:
    text = re.sub(r"<[^>]+>", " ", html_or_md)  # strip HTML tags
    text = re.sub(r"\s+", " ", text).strip()
    words = text.split()
    return " ".join(words[:max_words])

def normalize_links_to_style(text: str, style: str) -> str:
    md_link_re = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    html_link_re = re.compile(r'<a\s+[^>]*href=[\'"]([^\'"]+)[\'"][^>]*>(.*?)</a>', re.IGNORECASE)

    if style == "html":
        def md_to_html(m):
            return _wrap_link(m.group(1), m.group(2), "html")
        return md_link_re.sub(md_to_html, text)
    else:
        def html_to_md(m):
            inner = re.sub(r"<.*?>", "", m.group(2))
            return f"[{inner}]({m.group(1)})"
        return html_link_re.sub(html_to_md, text)

# ===== autolinker =====
STOPWORDS = {"and","or","the","a","an","for","with","to","of","in","on","at","by","from","your","my"}

def _plural_variants(word: str) -> set:
    v = {word}
    w = word
    if w.endswith("ies") and len(w) > 3:
        v.add(w[:-3] + "y")
    elif w.endswith("es") and len(w) > 2:
        v.add(w[:-2])
    elif w.endswith("s") and not w.endswith("ss"):
        v.add(w[:-1])
    else:
        v.add(w + "s")
    return v

def _ngrams(tokens: List[str], min_n=2, max_n=4) -> List[List[str]]:
    toks = [t for t in tokens if t and t not in STOPWORDS]
    out = []
    L = len(toks)
    for n in range(max_n, min_n - 1, -1):
        for i in range(L - n + 1):
            out.append(toks[i:i+n])
    return out

def load_sitemap_from_text(txt: str) -> List[str]:
    urls = []
    for line in txt.splitlines():
        u = line.strip()
        if u.startswith(("http://","https://")):
            urls.append(u)
    seen=set(); out=[]
    for u in urls:
        if u not in seen:
            out.append(u); seen.add(u)
    return out

def filter_candidate_urls(urls: List[str]) -> List[str]:
    drops = ("/privacy","/terms","/disclaimer","/contact","/about","/category/","/features/","/instagram")
    keep=[]
    for u in urls:
        low=u.lower().rstrip("/")
        if low in {"https://www.tastetorate.com","http://www.tastetorate.com"}:
            continue
        if any(d in low for d in drops):
            continue
        last = low.split("/")[-1]
        if last:
            keep.append(u)
    return keep

def build_link_index(urls: List[str]) -> List[dict]:
    index=[]
    for u in urls:
        slug = u.rstrip("/").split("/")[-1]
        raw_tokens = re.split(r"[-_]+", slug.lower())
        raw_tokens = [t for t in raw_tokens if t and t not in STOPWORDS]
        base_tokens = [t for t in raw_tokens if t not in {"recipe","recipes"}]
        for toks in _ngrams(base_tokens, min_n=2, max_n=4):
            last_variants = _plural_variants(toks[-1])
            for last in last_variants:
                phrase_tokens = toks[:-1] + [last]
                phrase = " ".join(phrase_tokens)
                if len(phrase) < 6:
                    continue
                pattern = re.compile(rf"(?i)(?<!\])\b{re.escape(phrase)}\b(?!\()")
                index.append({"phrase": phrase, "url": u, "pattern": pattern})
        full_title = slug_to_title(slug)
        if len(full_title.split()) >= 2:
            pattern = re.compile(rf"(?i)(?<!\>)\b{re.escape(full_title)}\b(?!\()")
            index.append({"phrase": full_title, "url": u, "pattern": pattern})
    index.sort(key=lambda d: len(d["phrase"]), reverse=True)
    return index

def _wrap_link(text: str, url: str, style: str) -> str:
    if style == "html":
        return f'<a href="{url}" target="_blank" rel="nofollow noopener">{text}</a>'
    return f"[{text}]({url})"

def autolink_content(
    content: str,
    index: List[dict],
    max_links: int = 12,
    per_paragraph_max: int = 3,
    link_headings: bool = False,
    skip_lines_with_links: bool = True,
    link_style: str = "html",
) -> str:
    if not index or max_links <= 0:
        return content

    # Split content into sections based on H2 headings for better link distribution
    lines = content.splitlines()
    sections = []
    current_section = []
    
    for line in lines:
        if line.strip().startswith("## ") and current_section:
            sections.append(current_section)
            current_section = [line]
        else:
            current_section.append(line)
    
    if current_section:
        sections.append(current_section)
    
    # If no H2 sections found, process as single section
    if len(sections) <= 1:
        return _autolink_section(content, index, max_links, per_paragraph_max, link_headings, skip_lines_with_links, link_style)
    
    # Distribute links evenly across sections (minimum 2 links per section if possible)
    links_per_section = max(2, max_links // len(sections))
    remaining_links = max_links
    used_urls = set()
    processed_sections = []
    
    for i, section_lines in enumerate(sections):
        section_content = "\n".join(section_lines)
        # For last section, use all remaining links
        section_max_links = remaining_links if i == len(sections) - 1 else min(links_per_section, remaining_links)
        
        if section_max_links > 0:
            # Filter out already used URLs
            available_index = [item for item in index if item["url"] not in used_urls]
            processed_section = _autolink_section(section_content, available_index, section_max_links, per_paragraph_max, link_headings, skip_lines_with_links, link_style)
            
            # Track which URLs were used in this section
            for item in available_index[:section_max_links]:
                if item["pattern"].search(processed_section):
                    used_urls.add(item["url"])
                    remaining_links -= 1
                    if remaining_links <= 0:
                        break
        else:
            processed_section = section_content
            
        processed_sections.append(processed_section)
    
    return "\n\n".join(processed_sections)

def _autolink_section(
    content: str,
    index: List[dict],
    max_links: int = 12,
    per_paragraph_max: int = 3,
    link_headings: bool = False,
    skip_lines_with_links: bool = True,
    link_style: str = "html",
) -> str:
    """Helper function to process a single section with autolinking"""
    if not index or max_links <= 0:
        return content

    used_urls = set()
    total = 0
    out_lines = []
    para_count = 0

    md_link_pat = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    html_link_pat = re.compile(r'<a\s+[^>]*href=[\'"][^\'"]+[\'"][^>]*>.*?</a>', re.IGNORECASE)

    def _flush_paragraph():
        nonlocal para_count
        para_count = 0

    for raw in content.splitlines():
        line = raw
        stripped = line.lstrip()

        if not stripped:
            out_lines.append(line)
            _flush_paragraph()
            continue

        if stripped.startswith("#") and not link_headings:
            out_lines.append(line)
            continue

        if skip_lines_with_links and (md_link_pat.search(line) or html_link_pat.search(line)):
            out_lines.append(line)
            continue

        for item in index:
            if total >= max_links or para_count >= per_paragraph_max:
                break
            if item["url"] in used_urls:
                continue

            def _repl(m):
                nonlocal total, para_count
                if total >= max_links or para_count >= per_paragraph_max:
                    return m.group(0)
                used_urls.add(item["url"])
                total += 1
                para_count += 1
                return _wrap_link(m.group(0), item["url"], link_style)

            new_line, n = item["pattern"].subn(_repl, line, count=1)
            if n > 0:
                line = new_line
        out_lines.append(line)

    return "\n".join(out_lines)

# ---------------------------------------------------------
# DOCX export
# ---------------------------------------------------------
def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single')
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0000FF')
    rPr.append(u); rPr.append(color)
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def build_docx_from_content(topic: str, content: str) -> bytes:
    doc = docx.Document()
    doc.add_heading(f"Blog about {topic}", level=0)
    md_link_re = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
    html_link_re = re.compile(r'<a\s+[^>]*href=[\'"]([^\'"]+)[\'"][^>]*>(.*?)</a>', re.IGNORECASE)
    for raw_line in content.splitlines():
        p = doc.add_paragraph()
        line = raw_line
        matches = []
        for m in md_link_re.finditer(line): matches.append(("md", m.start(), m.end(), m))
        for m in html_link_re.finditer(line): matches.append(("html", m.start(), m.end(), m))
        matches.sort(key=lambda x: x[1])
        last_end = 0
        for kind, start, end, m in matches:
            before = line[last_end:start]
            if before: p.add_run(before)
            if kind == "md":
                text = m.group(1); url = m.group(2)
            else:
                url = m.group(1); text = re.sub(r"<.*?>", "", m.group(2))
            _add_hyperlink(p, url, text); last_end = end
        tail = line[last_end:]
        if tail: p.add_run(tail)
    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio.read()

# ---------------------------------------------------------
# Recipe parsing + fallback card
# ---------------------------------------------------------
def _to_iso8601_minutes(total_minutes: int) -> Optional[str]:
    if total_minutes <= 0:
        return None
    h = total_minutes // 60
    m = total_minutes % 60
    if h and m:
        return f"PT{h}H{m}M"
    if h:
        return f"PT{h}H"
    return f"PT{m}M"

def _extract_minutes(line: str) -> int:
    mins = 0
    h = re.search(r'(\d+)\s*h', line)
    m = re.search(r'(\d+)\s*m', line)
    if h: mins += int(h.group(1)) * 60
    if m: mins += int(m.group(1))
    if mins == 0:
        m2 = re.search(r'(\d+)', line)
        mins = int(m2.group(1)) if m2 else 0
    return mins

def parse_recipe_text_blocks(recipe_text: str) -> Dict[str, Any]:
    """Forgiving parser for the text block pasted in the UI."""
    lines = [l.rstrip() for l in recipe_text.splitlines()]
    title = ""
    description = ""
    ingredients: List[str] = []
    instructions: List[str] = []
    notes: List[str] = []
    nutrition: Dict[str, str] = {}
    category = None
    cuisine = None
    method = None
    diet = None
    keywords: List[str] = []
    servings = None
    yield_str = None
    prep_iso = cook_iso = total_iso = None

    nonempty = [l for l in lines if l.strip()]
    if nonempty:
        title = nonempty[0].strip()
    if len(nonempty) > 1:
        description = nonempty[1].strip()

    for ln in lines:
        L = ln.lower().strip()
        if L.startswith("prep time"):
            prep_iso = _to_iso8601_minutes(_extract_minutes(L))
        elif L.startswith("cook time"):
            cook_iso = _to_iso8601_minutes(_extract_minutes(L))
        elif L.startswith("total time"):
            total_iso = _to_iso8601_minutes(_extract_minutes(L))
        elif L.startswith("yield"):
            if ":" in ln:
                yield_str = ln.split(":",1)[-1].strip()
        elif L.startswith("servings"):
            m = re.search(r'(\d+)', ln)
            if m:
                servings = f"{m.group(1)} servings"
        elif L.startswith("category:"):
            category = ln.split(":",1)[-1].strip()
        elif L.startswith("method:"):
            method = ln.split(":",1)[-1].strip()
        elif L.startswith("cuisine:"):
            cuisine = ln.split(":",1)[-1].strip()
        elif L.startswith("diet:"):
            diet_line = ln.split(":",1)[-1].strip()
            if diet_line:
                diet = diet_line
                if diet_line.lower() not in [x.lower() for x in keywords]:
                    keywords.append(diet_line)
        elif L.startswith("keywords:"):
            kws = ln.split(":",1)[-1].strip()
            if kws:
                for kw in re.split(r"[,;]", kws):
                    k = kw.strip()
                    if k and k.lower() not in [x.lower() for x in keywords]:
                        keywords.append(k)

    section = None
    for ln in lines:
        t = ln.strip()
        if re.match(r"^ingredients\s*$", t, re.IGNORECASE):
            section = "ingredients"; continue
        if re.match(r"^instructions\s*$", t, re.IGNORECASE):
            section = "instructions"; continue
        if re.match(r"^notes?\s*$", t, re.IGNORECASE):
            section = "notes"; continue
        if re.match(r"^nutrition\s*$", t, re.IGNORECASE):
            section = "nutrition"; continue
        if not t:
            continue

        if section == "ingredients":
            if not re.match(r"^(prep|cook|total|servings|yield)\b", t, re.IGNORECASE):
                ingredients.append(t.lstrip("•- ").strip())
        elif section == "instructions":
            clean = re.sub(r'^\s*\d+\s*[\.\)\-]\s*', '', t)
            instructions.append(clean)
        elif section == "notes":
            notes.append(t)
        elif section == "nutrition":
            m = re.match(r'^([^:]+):\s*(.+)$', t)
            if m:
                nutrition[m.group(1).strip()] = m.group(2).strip()

    for token in re.split(r"[,\|/–—-]+", title):
        tok = token.strip().lower()
        if tok and len(tok) >= 3:
            keywords.append(tok)

    return {
        "title": title,
        "description": description,
        "ingredients": [i for i in ingredients if i],
        "instructions": [s for s in instructions if s],
        "notes": [n for n in notes if n],
        "nutrition": nutrition if nutrition else None,
        "category": category,
        "cuisine": cuisine,
        "method": method,
        "diet": diet,
        "keywords": keywords[:12],
        "servings": servings,
        "yield": yield_str or servings,
        "prepISO": prep_iso,
        "cookISO": cook_iso,
        "totalISO": total_iso,
    }

def _make_recipe_halal(recipe: Dict[str, Any]) -> Dict[str, Any]:
    """Return a copy of recipe with certain ingredients replaced by suitable alternatives.
    Does not add any explicit mention of dietary terms in text; simply swaps items in-place."""
    if not recipe:
        return recipe
    # Patterns cover common non-permitted items and typical variants; replacements are culinary equivalents.
    replacements = [
        (r"\b(pork|bacon|ham|prosciutto|pancetta)\b", "turkey bacon"),
        (r"\bpork\s+sausage\b", "turkey sausage"),
        (r"\bpepperoni\b", "beef pepperoni"),
        (r"\bsalami\b", "beef salami"),
        (r"\bchorizo\b", "turkey chorizo"),
        (r"\blard\b", "vegetable shortening"),
        (r"\bgelatin\b", "agar-agar powder"),
        (r"\bmirin\b", "rice vinegar + sugar"),
        (r"\b(?:red|white)?\s*wine\b", "grape juice + a splash of vinegar"),
        (r"\bsherry\b", "apple cider vinegar + apple juice"),
        (r"\b(beer|ale|lager)\b", "apple cider + a splash of vinegar"),
        (r"\brum\b", "molasses + apple juice"),
        (r"\bbrandy\b", "apple juice"),
        (r"\b(bourbon|whiskey|whisky)\b", "apple cider + vanilla bean paste"),
        (r"\bbacon grease\b", "olive oil"),
        (r"\bprosciutto\b", "turkey ham"),
        (r"\bvanilla extract\b", "vanilla bean paste"),
    ]
    def _repl_text(t: str) -> str:
        if not t:
            return t
        for pat, rep in replacements:
            t = re.sub(pat, rep, t, flags=re.IGNORECASE)
        return t
    new_recipe = dict(recipe)
    new_recipe["description"] = _repl_text(recipe.get("description", ""))
    new_recipe["ingredients"] = [_repl_text(x) for x in recipe.get("ingredients", [])]
    new_recipe["instructions"] = [_repl_text(x) for x in recipe.get("instructions", [])]
    raw_notes = recipe.get("notes", [])
    normalized_notes = []
    if isinstance(raw_notes, str):
        normalized_notes = [s.strip() for s in raw_notes.splitlines() if s.strip()]
    elif isinstance(raw_notes, (list, tuple)):
        normalized_notes = [str(x).strip() for x in raw_notes if str(x).strip()]
    else:
        if raw_notes:
            try:
                normalized_notes = [str(raw_notes).strip()]
            except Exception:
                normalized_notes = []
    new_recipe["notes"] = [_repl_text(x) for x in normalized_notes]
    new_recipe["keywords"] = [_repl_text(x) for x in recipe.get("keywords", [])]
    if isinstance(recipe.get("nutrition"), dict):
        new_recipe["nutrition"] = {k: _repl_text(v) for k, v in recipe["nutrition"].items()}
    return new_recipe


def _extract_recipe_from_article_md(md: str) -> Dict[str, Any]:
    """Heuristic extraction of recipe sections from article markdown, then parsed via parse_recipe_text_blocks."""
    if not md:
        return {}
    import re as _re

    def _find_section(md_text: str, names: list) -> str:
        for nm in names:
            m = _re.search(rf"(^|\n)#+\s*{nm}\s*\n(.*?)(?=\n#+\s|\Z)", md_text, flags=_re.IGNORECASE | _re.DOTALL)
            if m:
                return m.group(2).strip()
        return ""

    def _parse_list(block: str) -> list:
        items = []
        for line in block.splitlines():
            if _re.match(r"^\s*[-*]\s+", line):
                items.append(_re.sub(r"^\s*[-*]\s+", "", line).strip())
            elif _re.match(r"^\s*\d+[\.\)]\s+", line):
                items.append(_re.sub(r"^\s*\d+[\.\)]\s+", "", line).strip())
        return [i for i in items if i]

    # Title
    title_match = _re.search(r"^\s*#\s+(.+)$", md, flags=_re.MULTILINE)
    title = title_match.group(1).strip() if title_match else (st.session_state.get("seo_title") or st.session_state.get("topic") or "Recipe").strip()

    # Description: first paragraph after title
    desc = ""
    if title_match:
        start = title_match.end()
        next_h = _re.search(r"\n#{1,6}\s", md[start:], flags=_re.MULTILINE)
        block = md[start: start + next_h.start()] if next_h else md[start:]
        paras = [p.strip() for p in _re.split(r"\n\s*\n", block) if p.strip()]
        if paras:
            desc = paras[0]

    # Sections
    ingredients_block = _find_section(md, ["ingredients"])
    instructions_block = _find_section(md, ["instructions", "method", "directions"])
    notes_block = _find_section(md, ["notes", "note"]) or ""

    ingredients = _parse_list(ingredients_block)
    instructions = _parse_list(instructions_block)
    notes = _parse_list(notes_block) if notes_block else []

    # Meta
    prep = _re.search(r"prep\s*time\s*[:\-]\s*([^\n]+)", md, flags=_re.IGNORECASE)
    cook = _re.search(r"cook\s*time\s*[:\-]\s*([^\n]+)", md, flags=_re.IGNORECASE)
    total = _re.search(r"total\s*time\s*[:\-]\s*([^\n]+)", md, flags=_re.IGNORECASE)
    yld = _re.search(r"(yield|servings?)\s*[:\-]\s*([^\n]+)", md, flags=_re.IGNORECASE)

    # Build a minimal normalized text to feed into existing forgiving parser
    lines = [title, desc]
    if prep: lines.append(f"Prep Time: {prep.group(1).strip()}")
    if cook: lines.append(f"Cook Time: {cook.group(1).strip()}")
    if total: lines.append(f"Total Time: {total.group(1).strip()}")
    if yld: lines.append(f"Yield: {yld.group(2).strip()}")

    if ingredients:
        lines.append("Ingredients")
        lines.extend(ingredients)
    if instructions:
        lines.append("Instructions")
        lines.extend(instructions)
    if notes:
        lines.append("Notes")
        lines.extend(notes)

    normalized_text = "\n".join(lines)
    try:
        return parse_recipe_text_blocks(normalized_text)
    except Exception:
        return {}


def generate_js_recipe_card(recipe: Dict[str, Any]) -> str:
    """Generate pure JavaScript fillRecipeForm() function for Tasty Recipe CPT integration."""
    import json
    
    def js_escape(text: str) -> str:
        """Escape text for JavaScript string literal."""
        if not text:
            return ""
        return text.replace("\\", "\\\\").replace("'", "\\'").replace("\n", "\\n").replace("\r", "")
    
    def format_list_for_js(items: List[str], formatter: str = "p") -> str:
        """Format list items for JavaScript template literal.
        Accepts list/tuple or a single string; splits multi-line strings into items and strips bullet chars.
        """
        if items is None:
            return ""
        # Normalize input to a flat list of strings (not characters)
        norm_items: List[str] = []
        if isinstance(items, str):
            candidates = items.splitlines()
        elif isinstance(items, (list, tuple)):
            candidates = []
            for it in items:
                if isinstance(it, str):
                    candidates.extend(it.splitlines())
                else:
                    candidates.append(str(it))
        else:
            candidates = [str(items)]
        for p in candidates:
            p = p.strip("•- \t")
            if p:
                norm_items.append(p)
        escaped_items = [js_escape(item) for item in norm_items]
        # Return plain newline-separated items; the JS p()/n() helpers will wrap/number them in the editor.
        return "\\n".join(escaped_items)
    
    # Extract and format recipe data
    # Title/Author with safe defaults from session state if available
    try:
        _seo_title = st.session_state.get("seo_title")
        _focus_kw = st.session_state.get("focus_keyword")
        _author_name = st.session_state.get("author_name")
    except Exception:
        _seo_title = _focus_kw = _author_name = ""
    title = js_escape(recipe.get("title") or _seo_title or _focus_kw or "")
    author = js_escape(recipe.get("author") or _author_name or "")
    
    description = js_escape(recipe.get("description", ""))
    ingredients = format_list_for_js(recipe.get("ingredients", []), "p")
    instructions = format_list_for_js(recipe.get("instructions", []), "n")
    notes = format_list_for_js(recipe.get("notes", []), "p")
    
    # Extract times and convert from ISO to minutes (accept plain minutes too)
    def iso_to_minutes(iso_time) -> str:
        """Convert PT15M or PT1H30M to minutes; accept numeric/int or numeric string as minutes."""
        if iso_time is None:
            return ""
        try:
            import re
            if isinstance(iso_time, (int, float)):
                return str(int(iso_time))
            s = str(iso_time).strip()
            if not s:
                return ""
            if re.fullmatch(r"\d+", s):
                return s
            hours = re.search(r"(\d+)H", s)
            minutes = re.search(r"(\d+)M", s)
            total_minutes = 0
            if hours:
                total_minutes += int(hours.group(1)) * 60
            if minutes:
                total_minutes += int(minutes.group(1))
            return str(total_minutes) if total_minutes > 0 else ""
        except Exception:
            return ""
    
    # Prefer ISO fields, but also accept already-minutes fields if present
    prep_time = iso_to_minutes(recipe.get("prepISO", "") or recipe.get("prep_time", ""))
    cook_time = iso_to_minutes(recipe.get("cookISO", "") or recipe.get("cook_time", ""))
    total_time = iso_to_minutes(recipe.get("totalISO", "") or recipe.get("total_time", ""))
    
    # Extract other fields
    yield_value = js_escape(recipe.get("yield", ""))
    category = js_escape(recipe.get("category", ""))
    method = js_escape(recipe.get("method", ""))
    cuisine = js_escape(recipe.get("cuisine", ""))
    # Keywords: support list or comma-separated string
    kw_raw = recipe.get("keywords") or []
    if isinstance(kw_raw, list):
        keywords = js_escape(", ".join(kw_raw))
    else:
        keywords = js_escape(str(kw_raw))
    diet = js_escape(recipe.get("diet", ""))
    
    # Nutrition fields (support common key variants)
    nutrition = recipe.get("nutrition") or {}
    def nv(key: str, fallback_key: Optional[str] = None) -> str:
        v = nutrition.get(key)
        if (not v) and fallback_key:
            v = nutrition.get(fallback_key)
        return js_escape(str(v)) if v else ""
    serving_size = nv("Serving Size", "serving_size")
    calories = nv("Calories", "calories")
    sugar = nv("Sugar", "sugar")
    sodium = nv("Sodium", "sodium")
    fat = nv("Total Fat", "Fat")
    saturated_fat = nv("Saturated Fat", "saturated_fat")
    unsaturated_fat = nv("Unsaturated Fat", "unsaturated_fat")
    trans_fat = nv("Trans Fat", "trans_fat")
    cholesterol = nv("Cholesterol", "cholesterol")
    carbohydrates = nv("Total Carbohydrates", "Carbohydrates")
    fiber = nv("Dietary Fiber", "Fiber")
    protein = nv("Protein", "protein")
    
    # Apply minimal derivations; avoid fabricating defaults
    # Derive total_time if missing by summing provided prep and cook times
    try:
        pt_int = int(prep_time) if (prep_time and prep_time.isdigit()) else 0
        ct_int = int(cook_time) if (cook_time and cook_time.isdigit()) else 0
    except Exception:
        pt_int, ct_int = 0, 0
    if (not total_time or not total_time.isdigit()) and (pt_int or ct_int):
        total_time = str(pt_int + ct_int)
    # If still no prep_time but total_time present (and no cook time), use total as prep
    if (not prep_time) and (total_time and total_time.isdigit()) and (not cook_time or not cook_time.isdigit()):
        prep_time = total_time

    # Yield: only use explicit values from recipe; do not fabricate defaults
    if not yield_value:
        y = recipe.get("yield", "") or recipe.get("servings", "")
        yield_value = js_escape(y)

    # Category: do not infer; leave blank if not provided
    title_txt = (recipe.get("title", "") or "")
    title_l = title_txt.lower()
    kw_list = recipe.get("keywords", []) or []
    hay = (" ".join(kw_list) + " " + title_l).strip()
    if not category:
        pass

    # Method: leave blank if not provided
    # (no default)

    # Cuisine: leave blank if not provided
    # (no default)

    # Keywords: leave blank if not provided
    # (no inference from title)

    # Diet: leave blank if not provided
    # (no default)

    # Nutrition: use only values from context; do not fabricate defaults
    # Leave serving_size, calories, sugar, sodium, fat, saturated_fat, unsaturated_fat,
    # trans_fat, cholesterol, carbohydrates, fiber, protein as-is.

    # Final pass: derive times when supported by provided values; do not force zeros
    try:
        _pt = int(prep_time) if (prep_time and prep_time.isdigit()) else None
        _ct = int(cook_time) if (cook_time and cook_time.isdigit()) else None
        _tt = int(total_time) if (total_time and total_time.isdigit()) else None
    except Exception:
        _pt = _ct = _tt = None
    if (_tt is None) and ((_pt is not None) or (_ct is not None)):
        total_time = str(((_pt or 0) + (_ct or 0))) if ((_pt or 0) + (_ct or 0)) > 0 else ""
    if (not cook_time) and (_tt is not None) and (_pt is not None) and (_tt - _pt >= 0):
        cook_time = str(_tt - _pt)
    # Do not set prep_time/cook_time/total_time to '0' when unknown

    # Generate the JavaScript function
    js_code = f"""function fillRecipeForm() {{
const f=(s,v)=>document.querySelector(s)&&(document.querySelector(s).value=v),
p=t=>t.split('\\n').map(x=>`<p>${{x}}</p>`).join(''),
n=t=>t.split('\\n').map((x,i)=>`<p>${{i+1}}. ${{x}}</p>`).join(''),
g={{
description:'{description}',
ingredients:p('{ingredients}'),
instructions:n('{instructions}'),
notes:p('{notes}')
}};
for(const k in g){{
const ed=window.tinyMCE?.get(`tasty-recipes-recipe-${{k}}`);
ed&&ed.setContent(g[k]);
}}
["prep_time","cook_time","total_time","yield","category","method","cuisine","keywords","diet",
"serving_size","calories","sugar","sodium","fat","saturated_fat","unsaturated_fat","trans_fat",
"cholesterol","carbohydrates","fiber","protein"
].forEach(k=>f(`[name="${{k}}"]`,{{
prep_time:'{prep_time}',
cook_time:'{cook_time}',
total_time:'{total_time}',
yield:'{yield_value}',
category:'{category}',
method:'{method}',
cuisine:'{cuisine}',
keywords:'{keywords}',
diet:'{diet}',
serving_size:'{serving_size}',
calories:'{calories}',
sugar:'{sugar}',
sodium:'{sodium}',
fat:'{fat}',
saturated_fat:'{saturated_fat}',
unsaturated_fat:'{unsaturated_fat}',
trans_fat:'{trans_fat}',
cholesterol:'{cholesterol}',
carbohydrates:'{carbohydrates}',
fiber:'{fiber}',
protein:'{protein}'
}}[k]));
}}
fillRecipeForm();"""
    
    return js_code

def synthesize_recipe_from_context(
    topic: str = "",
    focus_keyword: str = "",
    full_recipe_text: str = "",
    article_md: str = "",
    model: Optional[str] = None,
    temperature: Optional[float] = None,
) -> Dict[str, Any]:
    """
    Use the LLM to synthesize a well-organized recipe text from context, then parse it
    with parse_recipe_text_blocks. Uses the "Full Recipe (parsed for Tasty + fallback)"
    as inspiration when provided; otherwise falls back to the generated article content.
    Returns an empty dict on failure.
    """
    try:
        mdl = model or st.session_state.get("model_name", "gpt-4.1")
        temp = float(temperature if temperature is not None else st.session_state.get("temperature", 0.2))

        context_bits: List[str] = []
        if full_recipe_text and full_recipe_text.strip():
            context_bits.append("Full Recipe text provided by user:\n" + full_recipe_text.strip()[:5000])
        if article_md and article_md.strip():
            context_bits.append("Article content (markdown):\n" + article_md.strip()[:5000])
        context = "\n\n".join(context_bits) or "No explicit recipe text was provided."

        prompt = f"""
From the context below, synthesize a single, clean, well-organized recipe in PLAIN TEXT using EXACTLY this format so a simple parser can read it.

Title on the first line
Optional short description on the second line
Prep time: <number> m or h m
Cook time: <number> m or h m
Total time: <number> m or h m
Yield: <e.g., 12 cookies or 1 loaf>
Servings: <e.g., 12 servings>
Category: <e.g., Dessert>
Method: <e.g., Baked>
Cuisine: <e.g., American>
Diet: <e.g., Vegetarian>
Keywords: <comma-separated keywords>

Ingredients
- item 1
- item 2
- ...

Instructions
1. step 1
2. step 2
3. ...

Notes
- note line 1
- note line 2

Nutrition
Serving Size: <e.g., 1 cookie>
Calories: <e.g., 210 kcal>
Carbohydrates: <e.g., 28 g>
Protein: <e.g., 3 g>
Fat: <e.g., 9 g>
Saturated Fat: <e.g., 5 g>
Unsaturated Fat: <e.g., 3 g>
Trans Fat: <e.g., 0 g>
Cholesterol: <e.g., 20 mg>
Sodium: <e.g., 120 mg>
Fiber: <e.g., 1 g>
Sugar: <e.g., 16 g>

Rules:
- Use information from the provided context. When values are not explicitly given, infer reasonable, specific values based on the recipe's ingredients and method (do not leave items blank).
- If the context is insufficient to produce a specific recipe with concrete ingredients and steps, respond with the single word INSUFFICIENT.
- Use US measurements and kitchen conventions where appropriate; do not add commentary, ranges, or approximation words.
- Avoid any code fences, Markdown headings (#), or extra labeling not shown in the format.
- Keep the title concise and descriptive.
- Keep ingredient lines one per line; keep steps concise but precise.
- Provide nutrition values per serving with units (kcal, g, mg). Always include all fields above.

Context:
Topic: {topic}
Focus keyword: {focus_keyword}

{context}
""".strip()

        raw = _openai_text(prompt, mdl, temp)
        if raw and raw.strip().upper().startswith("INSUFFICIENT"):
            return {}
        parsed = parse_recipe_text_blocks(raw)
        if parsed and (parsed.get("ingredients") or parsed.get("instructions")):
            return parsed
        return {}
    except Exception:
        return {}

def html_recipe_fallback(recipe: Dict[str, Any]) -> str:
    """Generate a standardized recipe card using the brand template."""
    try:
        # Load the template and config
        template_path = Path(__file__).parent / "brand_assets" / "Recipe_Templates" / "recipe_card_template.html"
        config_path = Path(__file__).parent / "brand_assets" / "Recipe_Templates" / "recipe_config.json"
        
        if not template_path.exists():
            return _legacy_recipe_fallback(recipe)
            
        with open(template_path, 'r', encoding='utf-8') as f:
            template = f.read()
            
        # Load config for defaults
        config = {}
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
        
        defaults = config.get('recipe_defaults', {})
        
        # Prepare template variables
        recipe_title = recipe.get('title', 'Delicious Recipe')
        recipe_subtitle = recipe.get('subtitle', defaults.get('recipe_subtitle', 'Amazing Cozy Drink'))
        rating_stars = recipe.get('rating', defaults.get('rating_stars', '★★★★★'))
        author = recipe.get('author', defaults.get('author', 'Abby Martin'))
        
        # Format time information
        total_time = recipe.get('totalISO', defaults.get('default_total_time', '5 minutes'))
        servings = recipe.get('yield', defaults.get('default_servings', '1 serving'))
        diet_type = recipe.get('diet', defaults.get('diet_type', 'Vegetarian'))
        
        # Handle recipe image
        recipe_image_html = ''
        if recipe.get('image_url'):
            alt_text = recipe.get('image_alt', f'{recipe_title} recipe')
            recipe_image_html = f'<img class="recipe-image" src="{recipe["image_url"]}" alt="{alt_text}">'
        
        # Format description
        recipe_description = recipe.get('description', f'A healthy, quality {recipe_title.lower()} recipe made with green tea powder, steamed milk, and a touch of honey.')
        
        # Format ingredients list
        ingredients_list = ''
        if recipe.get('ingredients'):
            ingredients_list = '\n'.join([f'<li>{ingredient}</li>' for ingredient in recipe['ingredients']])
        
        # Format instructions list
        instructions_list = ''
        if recipe.get('instructions'):
            instructions_list = '\n'.join([f'<li>{instruction}</li>' for instruction in recipe['instructions']])
        
        # Format tips section
        tips_section = ''
        if recipe.get('notes'):
            tips_html = '\n'.join([f'<div class="tip-item">• {note}</div>' for note in recipe['notes']])
            tips_section = f'''
            <div class="recipe-tips">
                <div class="tips-title">Chef Tips</div>
                {tips_html}
            </div>'''
        
        # Format nutrition section
        nutrition_section = ''
        if recipe.get('nutrition'):
            nutrition_items = []
            for key, value in recipe['nutrition'].items():
                nutrition_items.append(f'<div class="nutrition-item"><span class="nutrition-value">{value}</span><br>{key}</div>')
            nutrition_grid = '\n'.join(nutrition_items)
            nutrition_section = f'''
            <div class="recipe-nutrition">
                <div class="nutrition-title">Nutrition</div>
                <div class="nutrition-grid">
                    {nutrition_grid}
                </div>
            </div>'''
        
        # Replace template placeholders
        html_output = template.format(
            recipe_title=recipe_title,
            recipe_subtitle=recipe_subtitle,
            rating_stars=rating_stars,
            author=author,
            total_time=total_time,
            servings=servings,
            diet_type=diet_type,
            recipe_image_html=recipe_image_html,
            recipe_description=recipe_description,
            ingredients_list=ingredients_list,
            instructions_list=instructions_list,
            tips_section=tips_section,
            nutrition_section=nutrition_section
        )
        
        return html_output
        
    except Exception as e:
        # Fallback to legacy format if template fails
        return _legacy_recipe_fallback(recipe)

def _legacy_recipe_fallback(recipe: Dict[str, Any]) -> str:
    """Legacy recipe card format as fallback."""
    parts = []
    parts.append('<div class="ttr-recipe-card">')
    parts.append(f'<h2 class="ttr-recipe-title">{recipe.get("title","Recipe")}</h2>')
    if recipe.get("description"):
        parts.append(f'<p class="ttr-recipe-summary">{recipe["description"]}</p>')
    times_bits = []
    if recipe.get("prepISO"): times_bits.append(f'Prep: {recipe["prepISO"]}')
    if recipe.get("cookISO"): times_bits.append(f'Cook: {recipe["cookISO"]}')
    if recipe.get("totalISO"): times_bits.append(f'Total: {recipe["totalISO"]}')
    if recipe.get("yield"): times_bits.append(f'Yield: {recipe["yield"]}')
    if times_bits:
        parts.append('<p class="ttr-recipe-times">' + " · ".join(times_bits) + "</p>")
    if recipe.get("ingredients"):
        parts.append("<h3>Ingredients</h3><ul>")
        for i in recipe["ingredients"]:
            parts.append(f"<li>{i}</li>")
        parts.append("</ul>")
    if recipe.get("instructions"):
        parts.append("<h3>Instructions</h3><ol>")
        for s in recipe["instructions"]:
            parts.append(f"<li>{s}</li>")
        parts.append("</ol>")
    if recipe.get("notes"):
        parts.append("<h3>Notes</h3><ul>")
        for n in recipe["notes"]:
            parts.append(f"<li>{n}</li>")
        parts.append("</ul>")
    if recipe.get("nutrition"):
        parts.append("<h3>Nutrition</h3><ul>")
        for k, v in recipe["nutrition"].items():
            parts.append(f"<li><strong>{k}:</strong> {v}</li>")
        parts.append("</ul>")
    parts.append("</div>")
    return "\n".join(parts)

# ---------------------------------------------------------
# WordPress helpers
# ---------------------------------------------------------
def clean_site_url(site_url: str) -> str:
    """Clean site URL by removing common WordPress admin paths."""
    url = site_url.strip()
    # Remove common WordPress admin paths
    suffixes_to_remove = ["/wp-admin", "/wp-admin/", "/wp-login.php"]
    for suffix in suffixes_to_remove:
        if url.endswith(suffix):
            url = url[:-len(suffix)]
    return url.rstrip("/")

def wp_base(site_url: str) -> str:
    return clean_site_url(site_url) + "/wp-json/wp/v2"

def wp_auth(username: str, app_password: str):
    return HTTPBasicAuth(username, app_password)

def test_wp_connection(site_url: str, username: str, app_password: str) -> dict:
    """Test WordPress REST API connection and return status info."""
    try:
        # Test basic API endpoint
        base_url = wp_base(site_url)
        r = requests.get(base_url, auth=wp_auth(username, app_password), timeout=30)
        
        if r.status_code == 200:
            return {"success": True, "message": "WordPress REST API connection successful", "api_info": r.json()}
        elif r.status_code == 401:
            return {"success": False, "message": "Authentication failed. Check username and application password."}
        elif r.status_code == 404:
            return {"success": False, "message": f"WordPress REST API not found at {base_url}. Check if site URL is correct and REST API is enabled."}
        else:
            response_text = r.text[:500]
            if response_text.strip().startswith('<!doctype html>') or response_text.strip().startswith('<html'):
                return {"success": False, "message": f"Received HTML page instead of API response. WordPress REST API may not be enabled or accessible at {site_url}"}
            else:
                return {"success": False, "message": f"API connection failed ({r.status_code}): {response_text}"}
    except requests.exceptions.RequestException as e:
        return {"success": False, "message": f"Network error: {str(e)}"}
    except Exception as e:
        return {"success": False, "message": f"Unexpected error: {str(e)}"}

def discover_tasty_recipe_cpt(site_url: str, username: str, app_password: str):
    try:
        url = site_url.rstrip("/") + "/wp-json/wp/v2/types"
        r = requests.get(url, auth=wp_auth(username, app_password), timeout=30)
        if r.status_code != 200:
            return None
        data = r.json()
        for cpt_slug, spec in data.items():
            low = (cpt_slug or "").lower()
            if "tasty" in low and "recipe" in low:
                rest_base = spec.get("rest_base") or cpt_slug
                return {"rest_base": rest_base, "slug": cpt_slug}
        return None
    except Exception:
        return None

def create_tasty_recipe_via_rest(site_url: str, username: str, app_password: str, recipe: dict) -> Optional[int]:
    info = discover_tasty_recipe_cpt(site_url, username, app_password)
    if not info:
        return None

    ingredients_html = ""
    if recipe.get("ingredients"):
        ingredients_html = "<h3>Ingredients</h3><ul>" + "".join([f"<li>{i}</li>" for i in recipe["ingredients"]]) + "</ul>"

    instructions_html = ""
    if recipe.get("instructions"):
        instructions_html = "<h3>Instructions</h3><ol>" + "".join([f"<li>{s}</li>" for s in recipe["instructions"]]) + "</ol>"

    notes_html = ""
    if recipe.get("notes"):
        notes_html = "<h3>Notes</h3><ul>" + "".join([f"<li>{n}</li>" for n in recipe["notes"]]) + "</ul>"

    times = []
    if recipe.get("prepISO"): times.append(f'Prep: {recipe["prepISO"]}')
    if recipe.get("cookISO"): times.append(f'Cook: {recipe["cookISO"]}')
    if recipe.get("totalISO"): times.append(f'Total: {recipe["totalISO"]}')
    if recipe.get("yield"): times.append(f'Yield: {recipe["yield"]}')
    times_html = f'<p>{" · ".join(times)}</p>' if times else ""

    inner = "".join([
        f'<p>{recipe.get("description","").strip()}</p>' if recipe.get("description") else "",
        times_html,
        ingredients_html,
        instructions_html,
        notes_html,
    ])

    payload = {"title": recipe.get("title") or "Recipe", "status": "publish", "content": inner}

    try:
        url = site_url.rstrip("/") + f"/wp-json/wp/v2/{info['rest_base']}"
        r = requests.post(url, json=payload, auth=wp_auth(username, app_password), timeout=60)
        if r.status_code in (200, 201):
            return r.json().get("id")
        return None
    except Exception:
        return None

def embed_tasty_recipe_shortcode(recipe_id: int) -> str:
    return f'<!-- wp:shortcode -->[tasty-recipe id="{recipe_id}"]<!-- /wp:shortcode -->'

def wp_find_term(site_url: str, username: str, app_password: str, taxonomy: str, name: str) -> Optional[int]:
    r = requests.get(wp_base(site_url) + f"/{taxonomy}", params={"search": name, "per_page": 100},
                     auth=wp_auth(username, app_password), timeout=30)
    if r.status_code == 200:
        for t in r.json():
            if t.get("name","").lower() == name.lower():
                return t.get("id")
    return None

def wp_create_term(site_url: str, username: str, app_password: str, taxonomy: str, name: str) -> Optional[int]:
    r = requests.post(wp_base(site_url) + f"/{taxonomy}", json={"name": name},
                      auth=wp_auth(username, app_password), timeout=30)
    if r.status_code in (200, 201):
        return r.json().get("id")
    return None

def ensure_terms(site_url: str, username: str, app_password: str, taxonomy: str, names: List[str]) -> List[int]:
    ids=[]
    for name in [n.strip() for n in (names or []) if n.strip()]:
        tid = wp_find_term(site_url, username, app_password, taxonomy, name)
        if not tid:
            tid = wp_create_term(site_url, username, app_password, taxonomy, name)
        if tid: ids.append(tid)
    return ids

def wp_upload_media(site_url: str, username: str, app_password: str,
                    file_bytes: bytes, filename: str, mime: str = "image/jpeg") -> int:
    files = {"file": (filename, file_bytes, mime)}
    r = requests.post(wp_base(site_url) + "/media", files=files,
                      auth=wp_auth(username, app_password), timeout=60)
    if r.status_code not in (200, 201):
        raise RuntimeError(f"Media upload failed ({r.status_code}): {r.text[:300]}")
    return r.json().get("id")

def wp_upload_media_with_metadata(site_url: str, username: str, app_password: str,
                                 file_bytes: bytes, filename: str, seo_metadata: dict,
                                 mime: str = "image/jpeg") -> dict:
    """
    Upload media to WordPress with comprehensive SEO metadata.
    Returns dict with media_id, url, and metadata info.
    """
    # Step 1: Upload the file
    files = {"file": (filename, file_bytes, mime)}
    r = requests.post(wp_base(site_url) + "/media", files=files,
                      auth=wp_auth(username, app_password), timeout=60)
    if r.status_code not in (200, 201):
        raise RuntimeError(f"Media upload failed ({r.status_code}): {r.text[:300]}")
    
    media_data = r.json()
    media_id = media_data.get("id")
    
    # Step 2: Update media with SEO metadata
    if media_id and seo_metadata:
        update_payload = {}
        
        # Map SEO metadata to WordPress fields
        if seo_metadata.get('alt_text'):
            update_payload['alt_text'] = seo_metadata['alt_text']
        
        if seo_metadata.get('caption'):
            update_payload['caption'] = seo_metadata['caption']
        
        if seo_metadata.get('description'):
            update_payload['description'] = seo_metadata['description']
        
        # Update the media item with metadata
        if update_payload:
            update_r = requests.post(
                wp_base(site_url) + f"/media/{media_id}",
                json=update_payload,
                auth=wp_auth(username, app_password),
                timeout=30
            )
            if update_r.status_code not in (200, 201):
                # Don't fail the entire upload if metadata update fails
                print(f"Warning: Metadata update failed for media {media_id}: {update_r.text[:200]}")
    
    return {
        'media_id': media_id,
        'url': media_data.get('source_url', ''),
        'filename': filename,
        'seo_metadata': seo_metadata
    }

def wp_create_post(site_url: str, username: str, app_password: str,
                   title: str, content_html: str, status: str = "draft",
                   category_ids: Optional[List[int]] = None, tag_ids: Optional[List[int]] = None,
                   featured_media_id: Optional[int] = None, excerpt: Optional[str] = None,
                   slug: Optional[str] = None) -> dict:
    payload = {"title": title, "content": content_html, "status": status}
    if category_ids: payload["categories"] = category_ids
    if tag_ids: payload["tags"] = tag_ids
    if featured_media_id: payload["featured_media"] = featured_media_id
    if excerpt: payload["excerpt"] = excerpt
    if slug: payload["slug"] = slug
    
    endpoint_url = wp_base(site_url) + "/posts"
    
    try:
        r = requests.post(endpoint_url, json=payload,
                          auth=wp_auth(username, app_password), timeout=60)
        
        if r.status_code not in (200, 201):
            # Check if response looks like HTML (website error page) vs JSON (API error)
            response_text = r.text[:1000]
            if response_text.strip().startswith('<!doctype html>') or response_text.strip().startswith('<html'):
                raise RuntimeError(f"Post create failed ({r.status_code}): WordPress REST API endpoint not found or not accessible. Check if:\n1. Site URL is correct: {site_url}\n2. WordPress REST API is enabled\n3. Application password has proper permissions\n4. Site is accessible\n\nReceived HTML error page instead of API response.")
            else:
                raise RuntimeError(f"Post create failed ({r.status_code}): {response_text[:500]}")
        
        return r.json()
    except requests.exceptions.RequestException as e:
        raise RuntimeError(f"Network error connecting to WordPress API at {endpoint_url}: {str(e)}")
    except Exception as e:
        raise RuntimeError(f"Unexpected error during post creation: {str(e)}")

# ---------------------------------------------------------
# SEO Title Generation
# ---------------------------------------------------------

def generate_seo_title(topic: str, focus_keyword: str, model: str = "gpt-4.1", temperature: float = 0.7) -> str:
    """
    Generate a creative, SEO-optimized title that naturally incorporates the focus keyword.
    """
    
    prompt = f"""
Create a highly engaging, SEO-optimized blog post title for a food blog article.

Topic: {topic}
Focus Keyword: {focus_keyword}

Requirements:
1. Include the focus keyword naturally and prominently
2. Keep it between 50-60 characters for optimal SEO
3. Make it click-worthy, engaging, and appealing to food lovers
4. Use creative, natural language that sounds authentic
5. Avoid overused words like "Ultimate", "Amazing", "Perfect" unless they truly fit
6. Focus on what makes this recipe/topic unique or special
7. Consider using numbers, questions, or intriguing phrases when appropriate

Examples of engaging food blog titles:
- "5-Minute Chocolate Mug Cake (No Mixer Required!)"
- "Why This Banana Bread Recipe Changed Everything"
- "Crispy Chicken Thighs: The Secret is in the Skin"
- "Homemade Pasta Sauce That Beats Store-Bought"
- "One-Bowl Brownies (Seriously Fudgy Results)"
- "The Only Pizza Dough Recipe You'll Ever Need"

Create a title that feels fresh, specific, and makes readers want to click. Generate ONE title only. Do not include quotes or explanations.
"""
    
    try:
        client = get_client()
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=100
        )
        title = response.choices[0].message.content.strip()
        # Remove quotes if present
        title = title.strip('"\'')
        return title
    except Exception as e:
        # Fallback to original topic if generation fails
        return topic

def generate_rankmath_seo_metadata(topic: str, focus_keyword: str, article_content: str, model: str = "gpt-4.1") -> dict:
    """
    Generate comprehensive RankMath SEO metadata including optimized titles, permalinks, 
    meta descriptions, focus keywords, schema markup, and social media snippets.
    """
    
    prompt = f"""
Generate comprehensive SEO metadata for a food blog article that will be optimized for RankMath SEO plugin.

Topic: {topic}
Focus Keyword: {focus_keyword}
Article Content Preview: {article_content[:500]}...

Generate the following SEO elements in JSON format with STRICT character limits:

1. SEO Title (EXACTLY 50-60 characters, include focus keyword naturally)
2. Meta Description (EXACTLY 150-160 characters, compelling and keyword-rich)
3. Permalink/Slug (MAXIMUM 75 characters, SEO-friendly URL slug)
4. Focus Keywords (primary + 2-3 secondary keywords)
5. Schema Markup (Recipe schema elements)
6. Open Graph Title (engaging for social media)
7. Open Graph Description (compelling social description)
8. Twitter Title (optimized for Twitter cards)
9. Twitter Description (Twitter-specific description)
10. Additional SEO Tags (relevant tags for better categorization)

CRITICAL REQUIREMENTS:
- SEO Title: Must be between 50-60 characters (optimal for search results)
- Meta Description: Must be between 150-160 characters (optimal for search snippets)
- Permalink Slug: Must be under 75 characters (optimal for URLs)
- All text should be engaging and natural
- Include focus keyword strategically but naturally
- Optimize for click-through rates
- Follow SEO best practices
- Make social media snippets compelling

Return ONLY a valid JSON object with these keys:
{{
  "seo_title": "...",
  "meta_description": "...",
  "permalink_slug": "...",
  "focus_keyword": "...",
  "secondary_keywords": ["...", "...", "..."],
  "schema_type": "Recipe",
  "schema_elements": {{
    "name": "...",
    "description": "...",
    "keywords": "...",
    "recipeCategory": "...",
    "recipeCuisine": "..."
  }},
  "og_title": "...",
  "og_description": "...",
  "twitter_title": "...",
  "twitter_description": "...",
  "additional_tags": ["...", "...", "..."]
}}
"""
    
    try:
        client = get_client()
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=1000
        )
        
        metadata_text = response.choices[0].message.content.strip()
        
        # Clean up the response to ensure it's valid JSON
        if metadata_text.startswith('```json'):
            metadata_text = metadata_text[7:]
        if metadata_text.endswith('```'):
            metadata_text = metadata_text[:-3]
        
        metadata = json.loads(metadata_text)
        
        # Add generated timestamp
        metadata['generated_at'] = datetime.datetime.now().isoformat()
        
        return metadata
        
    except Exception as e:
        # Fallback metadata if generation fails - respecting character limits
        fallback_title = f"{focus_keyword} - {topic}"
        if len(fallback_title) > 60:
            fallback_title = f"{focus_keyword} Recipe"[:60]
        elif len(fallback_title) < 50:
            fallback_title = f"{focus_keyword} - Easy {topic} Recipe"[:60]
            
        fallback_desc = f"Learn how to make {focus_keyword}. This {topic.lower()} recipe is easy, delicious, and perfect for any occasion. Get step-by-step instructions now!"
        if len(fallback_desc) > 160:
            fallback_desc = f"Learn how to make {focus_keyword}. This {topic.lower()} recipe is easy, delicious, and perfect for any occasion."[:160]
        elif len(fallback_desc) < 150:
            fallback_desc = f"Learn how to make {focus_keyword}. This {topic.lower()} recipe is easy, delicious, and perfect for any occasion. Try it today!"[:160]
            
        fallback_slug = to_slug(f"{focus_keyword}-{topic}")
        if len(fallback_slug) > 75:
            fallback_slug = to_slug(f"{focus_keyword}-recipe")[:75]
            
        return {
            "seo_title": fallback_title,
            "meta_description": fallback_desc,
            "permalink_slug": fallback_slug,
            "focus_keyword": focus_keyword,
            "secondary_keywords": ["recipe", "cooking", "homemade"],
            "schema_type": "Recipe",
            "schema_elements": {
                "name": topic,
                "description": f"A delicious {focus_keyword} recipe",
                "keywords": f"{focus_keyword}, recipe, cooking",
                "recipeCategory": "Main Course",
                "recipeCuisine": "American"
            },
            "og_title": f"{focus_keyword} Recipe",
            "og_description": f"Make this amazing {focus_keyword} with our easy recipe!",
            "twitter_title": f"{focus_keyword} Recipe",
            "twitter_description": f"Delicious {focus_keyword} recipe you'll love!",
            "additional_tags": ["recipe", "cooking", "food"],
            "generated_at": datetime.datetime.now().isoformat(),
            "fallback": True
        }

def generate_random_seed() -> int:
    """
    Generate a random seed for MidJourney consistency.
    """
    import random
    return random.randint(1000, 9999999)

def generate_midjourney_prompts(article_content: str, topic: str, focus_keyword: str = "", model: str = "gpt-4.1") -> dict:
    """
    Generate 7 MidJourney image prompts with comprehensive SEO metadata and placement information.
    """
    client = get_client()
    if not client:
        return {}
    
    # Generate a unique seed for this recipe batch
    seed = generate_random_seed()
    
    # Use focus keyword or extract from topic
    if not focus_keyword:
        focus_keyword = topic
    
    # Universal style anchor
    style_anchor = "Exact same batch as the featured image. Styled consistently in a bright Scandinavian-style kitchen with white marble countertops, soft natural window light, and minimal decor. Same props, same colors, same food batch, cinematic food photography style, shallow depth of field."
    
    prompt = f"""
You are a professional food photography director and SEO expert specializing in MidJourney prompts.

Analyze this recipe article and generate exactly 7 MidJourney image prompts with comprehensive SEO metadata:

Article Topic: {topic}
Focus Keyword: {focus_keyword}
Article Content:
{article_content[:3000]}...

Generate prompts for these 7 images in order:
1. Featured Image (Hero shot – top of article)
2. Ingredients Image (Immediately after the Ingredients section ends — not within instructions)
3. Step Image 1 (Place after Instruction step 1 — show only the action/technique in step 1)
4. Step Image 2 (Place after Instruction step 2 — show only the action/technique in step 2)
5. Step Image 3 (Place after Instruction step 3 — show only the action/technique in step 3)
6. Serving Image (Within the Serving/Serving Suggestions section if it exists)
7. Recipe Card Image (Clean top-down, end of article)

STRICT RULES FOR STEP IMAGES (step1, step2, step3):
- The prompt must focus ONLY on the instruction action and technique for that specific step.
- DO NOT include serving, plating, garnishing, ingredients flat-lay, or finished hero context.
- Use imperative, instructional language with concrete details (tools, cookware, angle, textures, doneness cues, hands-in-frame if relevant).
- Maintain exact continuity with the SAME batch as featured via the style anchor and the same seed.

For each image, provide:
- A detailed MidJourney prompt with style anchor and seed
- Exact placement location in the article
- Brief description of what the image shows
- Complete SEO metadata (alt text, filename, caption, description)

SEO Requirements:
- Alt Text: Must include exact keyword "{focus_keyword}"
- Filename: Hyphenated, lowercase, include keyword (e.g., {focus_keyword.lower().replace(' ', '-')}-step1.jpg)
- Caption: Short, descriptive, human-readable sentence
- Description: Full sentence describing dish with continuity reference

Use this seed for ALL prompts: {seed}
Include this style anchor in ALL prompts: "{style_anchor}"

Return the response in this exact JSON format:
{{
  "seed": {seed},
  "focus_keyword": "{focus_keyword}",
  "images": [
    {{
      "type": "featured",
      "prompt": "Photo-realistic food photography of [dish name], hero shot of the finished recipe with all key details visible. Exact batch reference for later steps. {style_anchor} --seed {seed}",
      "placement": "Top of article (before introduction)",
      "description": "Hero shot of the finished dish",
      "seo_metadata": {{
        "alt_text": "Alt text including exact keyword '{focus_keyword}'",
        "filename": "suggested-filename-with-keyword.jpg",
        "caption": "Short descriptive caption for humans",
        "description": "Full sentence description with dish reference"
      }}
    }},
    {{
      "type": "ingredients",
      "prompt": "Flat lay of all ingredients for [dish name], arranged neatly on the same marble countertop as the featured image. Exact same kitchen, same props, preparing for the featured batch. {style_anchor} --seed {seed}",
      "placement": "After ingredients section",
      "description": "Ingredients layout",
      "seo_metadata": {{
        "alt_text": "Ingredients for {focus_keyword} arranged on marble countertop",
        "filename": "ingredients-filename.jpg",
        "caption": "Fresh ingredients caption",
        "description": "Ingredients description with continuity"
      }}
    }}
    // ... continue for all 7 images with complete SEO metadata
  ]
}}

Make the prompts specific to the actual recipe content. Replace [dish name] and [preparation steps] with actual details from the article. Ensure all SEO metadata fields are properly populated with keyword-optimized content.
"""
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a professional food photography director. Generate detailed MidJourney prompts with exact placement metadata."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content.strip()
        
        # Try to parse JSON response
        try:
            import json
            # Extract JSON from response if it's wrapped in markdown
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()
            
            return json.loads(content)
        except json.JSONDecodeError:
            # Fallback: create basic structure with SEO metadata
            keyword_slug = focus_keyword.lower().replace(' ', '-')
            return {
                "seed": seed,
                "focus_keyword": focus_keyword,
                "images": [
                    {
                        "type": "featured",
                        "prompt": f"Photo-realistic food photography of {topic}, hero shot of the finished recipe. {style_anchor} --seed {seed}",
                        "placement": "Top of article",
                        "description": "Hero shot of the finished dish",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} finished dish on elegant plate",
                            "filename": f"{keyword_slug}-featured.jpg",
                            "caption": f"Delicious {focus_keyword} ready to serve",
                            "description": f"This stunning {focus_keyword} showcases the perfect balance of flavors and presentation."
                        }
                    },
                    {
                        "type": "ingredients",
                        "prompt": f"Flat lay of ingredients for {topic}, arranged beautifully. {style_anchor} --seed {seed}",
                        "placement": "After ingredients section",
                        "description": "Ingredients layout",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} ingredients arranged on marble countertop",
                            "filename": f"{keyword_slug}-ingredients.jpg",
                            "caption": f"Fresh ingredients for making {focus_keyword}",
                            "description": f"These carefully selected ingredients come together to create the perfect {focus_keyword}."
                        }
                    },
                    {
                        "type": "step1",
                        "prompt": f"Instruction-only process photo for step 1 of {focus_keyword}: hands actively performing the technique with the correct tool and vessel; emphasize textures and doneness cues (e.g., glossy, thickened, bubbling). Tight composition, 45° angle or overhead as appropriate, no plating or serving context. {style_anchor} --seed {seed}",
                        "placement": "In instructions section after step 1",
                        "description": "Instructional action for step 1 (tools, textures, technique only)",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} step 1 instructional process shot",
                            "filename": f"{keyword_slug}-step1.jpg",
                            "caption": f"{focus_keyword}: step 1 technique in action",
                            "description": f"Close-up process image showing only the technique for step 1 while making {focus_keyword}, focusing on tools and textures with no serving context."
                        }
                    },
                    {
                        "type": "step2",
                        "prompt": f"Instruction-only process photo for step 2 of {focus_keyword}: demonstrate the mid-process technique with hands-in-frame if relevant; include cookware, utensils, and texture progression. Keep composition clean, avoid garnish/finished plating. {style_anchor} --seed {seed}",
                        "placement": "In instructions section after step 2",
                        "description": "Instructional action for step 2 (tools, textures, technique only)",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} step 2 instructional process shot",
                            "filename": f"{keyword_slug}-step2.jpg",
                            "caption": f"{focus_keyword}: step 2 technique in action",
                            "description": f"Process image focused strictly on step 2 while making {focus_keyword}, highlighting technique and texture progression with no serving or ingredients flat-lay."
                        }
                    },
                    {
                        "type": "step3",
                        "prompt": f"Instruction-only process photo for step 3 of {focus_keyword}: final technique or finish cue (e.g., stir until emulsified, bake until golden), show close texture/doneness indicators. No hero plating, no garnish. {style_anchor} --seed {seed}",
                        "placement": "In instructions section after step 3",
                        "description": "Instructional action for step 3 (tools, textures, technique only)",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} step 3 instructional process shot",
                            "filename": f"{keyword_slug}-step3.jpg",
                            "caption": f"{focus_keyword}: step 3 technique in action",
                            "description": f"Process image focused strictly on step 3 while making {focus_keyword}, emphasizing final technique and clear doneness cues without any serving context."
                        }
                    },
                    {
                        "type": "serving",
                        "prompt": f"Serving {topic}, elegant plating. {style_anchor} --seed {seed}",
                        "placement": "Before serving section",
                        "description": "Dish being served",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} being served on beautiful dinnerware",
                            "filename": f"{keyword_slug}-serving.jpg",
                            "caption": f"Serving the finished {focus_keyword}",
                            "description": f"This beautifully plated {focus_keyword} is ready to impress your guests."
                        }
                    },
                    {
                        "type": "recipe_card",
                        "prompt": f"Recipe card layout for {topic}, styled presentation. {style_anchor} --seed {seed}",
                        "placement": "End of article",
                        "description": "Recipe card presentation",
                        "seo_metadata": {
                            "alt_text": f"{focus_keyword} recipe card with complete instructions",
                            "filename": f"{keyword_slug}-recipe-card.jpg",
                            "caption": f"Complete {focus_keyword} recipe card",
                            "description": f"This comprehensive recipe card makes it easy to recreate this delicious {focus_keyword}."
                        }
                    }
                ]
            }
        
    except Exception as e:
        st.error(f"Image prompt generation failed: {e}")
        return {}

# ---------------------------------------------------------
# OpenAI generation helpers
# ---------------------------------------------------------
BASE_SYSTEM = """
You are an expert food blogger (40+ years of home and pro kitchen experience).
Write in clean, human, helpful English with a mix of short and longer sentences.
No URLs. Use US numerals. Light, on-topic humor only. Active voice only.
No code fences. Use Markdown headings (#, ##, ###). Never repeat headings back-to-back.
"""

SECTION_RULES = """
Length: 2–3 paragraphs. Each paragraph 3–6 sentences (no one-line paragraphs).
SEO: Keep the focus keyword density ~0.6–0.8% naturally.
Style: Be specific, practical, and honest. Use concrete kitchen examples.
Bold important terms with ** ** sparingly.
"""

FAQ_RULES = """
Produce EXACTLY 4 Q&A items.
Each answer must be 4–6 sentences of practical guidance.
Do not include code fences, tables, or URLs.
Use concise, descriptive question wording.
"""

def list_identities() -> Dict[str, str]:
    base = Path(__file__).with_name("brand_assets") / "Identities"
    result: Dict[str, str] = {}
    try:
        if base.exists():
            for p in base.glob("*.json"):
                display = p.stem
                try:
                    obj = json.loads(p.read_text(encoding="utf-8", errors="ignore"))
                    display = obj.get("name") or display
                except Exception:
                    pass
                result[display] = str(p)
    except Exception:
        pass
    return result

def compose_system_instructions() -> str:
    base = BASE_SYSTEM.strip()
    ident = st.session_state.get("writer_identity_data")
    if not ident:
        return base
    parts: List[str] = []
    parts.append("Adopt the following writer identity and brand voice.")
    if ident.get("name"):
        parts.append(f"Name: {ident.get('name')}")
    if ident.get("brand"):
        parts.append(f"Brand: {ident.get('brand')}")
    po = ident.get("persona_overview", {}) or {}
    if any(po.get(k) for k in ("background","tone","audience")):
        parts.append(
            "Persona Overview: " + 
            "; ".join([
                f"background={po.get('background','')}",
                f"tone={po.get('tone','')}",
                f"audience={po.get('audience','')}"
            ])
        )
    vs = ident.get("voice_and_style", {}) or {}
    if vs:
        parts.append("Voice & Style: " + "; ".join([f"{k}: {v}" for k, v in vs.items() if v]))
    fmt = ident.get("formatting_conventions", {}) or {}
    if fmt:
        parts.append("Formatting Conventions: " + "; ".join([f"{k}: {v}" for k, v in fmt.items() if v]))
    struct = ident.get("article_structure", []) or []
    if struct:
        parts.append("Article Structure Preferences: " + " | ".join(struct))
    sig = ident.get("signature_examples", {}) or {}
    if sig:
        parts.append("Signature Style Examples: " + "; ".join([f"{k}: {v}" for k, v in sig.items() if v]))
    identity_block = "\n".join(parts)
    return base + "\n\n" + identity_block

def _openai_text(user_prompt: str, model: str, temperature: float = 0.6) -> str:
    client = get_client()
    resp = client.responses.create(
        model=model,
        instructions=compose_system_instructions(),
        input=user_prompt,
        temperature=temperature,
    )
    return resp.output_text.strip()

def build_part_prompt(part_title: str, topic: str, focus_keyword: str, recipe_text: str,
                      phrase_hints: Optional[List[str]], target_words: int, seo_title: Optional[str] = None) -> str:
    hints = ""
    if phrase_hints:
        hints = f"\nWhen natural, include several of these internal phrases: {', '.join(phrase_hints[:20])}\n"
    # Allocate target words roughly across 4 parts + FAQ (FAQ has its own target)
    per_part = max(220, min(550, target_words // 5))
    
    # Use SEO title if provided, otherwise use the original topic
    title_to_use = seo_title if seo_title else topic
    
    return f"""
Write the section **{part_title}** for a food blog article titled: "{title_to_use}".

Focus keyword: "{focus_keyword}"

Context recipe (may reference when relevant):
{recipe_text[:2000]}

Constraints:
{SECTION_RULES}

Formatting:
- Use H2 heading (##) for the section title
- Can include 0-2 internal links naturally within the content
- Follow standard Markdown formatting

Target ~{per_part} words for this section.
Do not include any other sections or meta text—only this section content with an appropriate Markdown heading and body.
{hints}
""".strip()

def build_faq_prompt(topic: str, focus_keyword: str, example_faq: str,
                     recipe_text: str, target_words: int, seo_title: Optional[str] = None) -> str:
    # Slightly larger budget for FAQ total
    faq_target = max(320, min(700, target_words // 5))
    
    # Use SEO title if provided, otherwise use the original topic
    title_to_use = seo_title if seo_title else topic
    
    return f"""
Create a **FAQ** section for the food blog article titled: "{title_to_use}".

Focus keyword: "{focus_keyword}"

If provided, align with the themes in this example FAQ (but rewrite fully and improve):
{example_faq or '(none provided)'}

Context recipe (for accuracy, if relevant):
{recipe_text[:2000]}

Constraints:
{FAQ_RULES}

Target total ~{faq_target} words across all 4 items.
Output format:
### FAQ
1. Question?
   Answer paragraph(s).
2. Question?
   Answer paragraph(s).
3. Question?
   Answer paragraph(s).
4. Question?
   Answer paragraph(s).
""".strip()

def gpt_article_single(topic: str, faq_text: str, recipe_text: str, focus_keyword: str,
                       model: str, phrase_hints: Optional[List[str]], temperature: float, seo_title: Optional[str] = None) -> str:
    # kept for compatibility with your earlier single-call path
    hints_block = ""
    if phrase_hints:
        hints_block = f"\n\nWhen natural, try to include several of these phrases: {', '.join(phrase_hints[:15])}\n"
    
    # Use SEO title if provided, otherwise use the original topic
    title_to_use = seo_title if seo_title else topic
    instructions = f"""
{compose_system_instructions()}

# Length & Paragraphing
Target around 1,500–2,200 characters total (minimum 1,500).
Each subsection should be 2 paragraphs. Each paragraph 3–5 sentences. Avoid one-sentence paragraphs.

Formatting:
- Structured Markdown with ##, ### headings.
- Do not put a subheading immediately after a heading.
- Do not include any code fences.
- Do not include an H1 title at the beginning

Structure:
- Part 1) Memory that inspired this treat & why it's special
- Part 2) How to make the flavor with the mix & ingredients
- Part 3) Make-ahead & storage
- Part 4) Best ingredients & party variations
- Part 5) FAQ (EXACTLY 4 Q/As)

Content & SEO:
- Every sentence must serve the article; avoid filler like "dive into".
- Offer honest comparisons with reasoning and real kitchen examples.
- Keep the focus keyword density around 0.6–0.8%.
- Each FAQ answer should be 4–6 sentences with helpful, detailed guidance.

Create a food blog article based on this content:
Topic: {topic}

Focus keyword for SEO: {focus_keyword}

Example FAQ:
{faq_text}

Full recipe:
{recipe_text}
{hints_block}
""".strip()
    client = get_client()
    resp = client.responses.create(model=model, instructions=instructions, input=topic, temperature=temperature)
    return resp.output_text.strip()

def build_conclusion_prompt(topic: str, focus_keyword: str, target_words: int, seo_title: Optional[str] = None) -> str:
    # Small budget for conclusion
    conclusion_target = max(150, min(300, target_words // 8))
    
    # Use SEO title if provided, otherwise use the original topic
    title_to_use = seo_title if seo_title else topic
    
    return f"""
Create a **Conclusion** section for the food blog article titled: "{title_to_use}".

Focus keyword: "{focus_keyword}"

Constraints:
- Write a warm, encouraging wrap-up that reinforces the joy of making this treat
- Mention how this recipe brings people together or creates special moments
- Include a gentle call-to-action encouraging readers to try the recipe
- End with social media follow encouragement (mention following for more recipes)
- Keep it personal and authentic to the food blogger voice
- Use the focus keyword naturally once if possible

Target ~{conclusion_target} words total.
Output format:
## Conclusion
[2-3 paragraphs of warm conclusion content]

For more delicious recipes like this, follow us on [Facebook](https://www.facebook.com/tastetorate) and [Pinterest](https://www.pinterest.com/tastetorate)!
""".strip()

def gpt_article_multi(topic: str, faq_text: str, recipe_text: str, focus_keyword: str,
                      model: str, phrase_hints: Optional[List[str]],
                      target_words: int, temperature: float, seo_title: Optional[str] = None) -> str:
    # Generate parts 1–4 independently
    parts_meta = [
        ("## The Memory Behind This Treat",),
        ("## How To Make It (Mix & Ingredients)",),
        ("## Make-Ahead & Storage",),
        ("## Best Ingredients & Party Variations",),
    ]
    out_sections: List[str] = []
    for (title,) in parts_meta:
        p = build_part_prompt(title, topic, focus_keyword, recipe_text, phrase_hints, target_words, seo_title)
        section = _openai_text(p, model=model, temperature=temperature)
        out_sections.append(section.strip())

    # FAQ (4 items)
    faq_prompt = build_faq_prompt(topic, focus_keyword, faq_text, recipe_text, target_words, seo_title)
    faq_section = _openai_text(faq_prompt, model=model, temperature=temperature)

    # Conclusion with social media links
    conclusion_prompt = build_conclusion_prompt(topic, focus_keyword, target_words, seo_title)
    conclusion_section = _openai_text(conclusion_prompt, model=model, temperature=temperature)

    # Assemble sections without adding duplicate H1 title
    content = "\n\n".join(out_sections + [faq_section.strip(), conclusion_section.strip()])
    return content.strip()

# ---------------------------------------------------------
# Sitemap fetcher (txt + xml + sitemap index)
# ---------------------------------------------------------
def fetch_sitemap_urls_any(url: str, timeout=20) -> List[str]:
    url = _normalize_url(url)
    r = requests.get(url, timeout=timeout)
    r.raise_for_status()
    ctype = (r.headers.get("content-type") or "").lower()

    collected: List[str] = []

    def _fetch_xml(xml_text: str):
        try:
            root = ET.fromstring(xml_text)
        except Exception:
            return
        tag = (root.tag or "").lower()
        if tag.endswith("sitemapindex"):
            for sm in root.findall(".//{*}sitemap/{*}loc"):
                loc = (sm.text or "").strip()
                if loc:
                    try:
                        sr = requests.get(loc, timeout=timeout)
                        sr.raise_for_status()
                        _fetch_xml(sr.text)
                    except Exception:
                        continue
        elif tag.endswith("urlset"):
            for node in root.findall(".//{*}url/{*}loc"):
                loc = (node.text or "").strip()
                if loc:
                    collected.append(loc)

    if "text/plain" in ctype or url.lower().endswith(".txt"):
        collected.extend(load_sitemap_from_text(r.text))
    else:
        _fetch_xml(r.text)

    return filter_candidate_urls(collected)

# ---------------------------------------------------------
# Streamlit UI
# ---------------------------------------------------------
st.set_page_config(page_title="Food Blog Article Generator", page_icon="🍿", layout="wide")
st.title("🍿 Food Blog Article Generator")
st.caption("Auto-internal-linking from your sitemap + CTA + one-click publish to WordPress. Long-form multi-request mode for bigger, better SEO articles.")

# Sidebar: API key
with st.sidebar:
    if not os.getenv("API_KEY"):
        st.info("Set API_KEY in a .env file next to app1.py (or export it in your shell). Create a .env file with a line: API_KEY=sk-... (the file is already in .gitignore).")

# Sidebar: Writer Identity
with st.sidebar:
    st.subheader("✍️ Writer Identity")
    ids = list_identities()
    names = sorted(ids.keys())
    if names:
        default_name = st.session_state.get("writer_identity_name") or names[0]
        if default_name not in names:
            default_name = names[0]
        sel = st.selectbox("Select writer identity", names, index=names.index(default_name))
        st.session_state["writer_identity_name"] = sel
        try:
            path = ids[sel]
            data = json.loads(Path(path).read_text(encoding="utf-8", errors="ignore"))
            st.session_state["writer_identity_data"] = data
            brand = data.get("brand") or ""
            st.caption(f"Using: {data.get('name','')}" + (f" — {brand}" if brand else ""))
        except Exception as e:
            st.warning(f"Failed to load identity: {e}")
            st.session_state["writer_identity_data"] = None
    else:
        st.caption("No identities found. Add JSON files under brand_assets/Identities.")

# Sidebar: Generation controls
with st.sidebar:
    st.subheader("🧠 Generation")
    model_name = st.text_input("Model", value=st.session_state.get("model_name", "gpt-4.1"))
    st.session_state["model_name"] = model_name
    temperature = st.slider("Creativity (temperature)", 0.0, 1.2, st.session_state.get("temperature", 0.6), 0.1)
    st.session_state["temperature"] = temperature
    use_multi_call = st.checkbox("Long form (multi-call)", value=st.session_state.get("use_multi_call", True))
    st.session_state["use_multi_call"] = use_multi_call
    target_words = st.slider("Target total words", 1200, 2800, st.session_state.get("target_words", 1800), 100)
    st.session_state["target_words"] = target_words

# Sidebar: Social links + CTA
with st.sidebar:
    st.subheader("🔗 Social links")
    fb_url = st.text_input("Facebook URL", value=st.session_state.get("fb_url", "https://www.facebook.com/YourPage"))
    st.session_state["fb_url"] = fb_url
    pin_url = st.text_input("Pinterest URL", value=st.session_state.get("pin_url", "https://www.pinterest.com/YourProfile"))
    st.session_state["pin_url"] = pin_url
    append_cta = st.checkbox("Append follower CTA at the end", value=st.session_state.get("append_cta", True))
    st.session_state["append_cta"] = append_cta

# Sidebar: Link style
with st.sidebar:
    st.subheader("🔗 Link Style")
    _stored_style = st.session_state.get("link_style", "html")
    link_style = st.radio("Choose link format", ["html", "markdown"], index=0 if _stored_style == "html" else 1)
    st.session_state["link_style"] = link_style

# Sidebar: Sitemap for Auto-Linking
with st.sidebar:
    st.subheader("🧭 Sitemap for Auto-Linking")
    st.caption("Upload a .txt list of URLs (one per line) or fetch from URL. Cached for reuse across topics.")
    smp = st.file_uploader("Sitemap .txt", type=["txt"])
    default_path = st.text_input("Or load from a local file path", value=st.session_state.get("default_path", ""))
    st.session_state["default_path"] = default_path
    sitemap_url = st.text_input("Or fetch from a sitemap URL (.txt or .xml)", value=st.session_state.get("sitemap_url", ""))
    st.session_state["sitemap_url"] = sitemap_url

    if "sitemap_urls" not in st.session_state:
        st.session_state["sitemap_urls"] = []

    c1, c2 = st.columns(2)
    if c1.button("Load sitemap"):
        try:
            if smp is not None:
                txt = smp.read().decode("utf-8", errors="ignore")
            elif default_path.strip():
                txt = Path(default_path.strip()).read_text(encoding="utf-8", errors="ignore")
            else:
                st.warning("Provide a sitemap file or a valid path first"); txt = ""
            st.session_state["sitemap_urls"] = load_sitemap_from_text(txt)
            st.success(f"Loaded {len(st.session_state['sitemap_urls'])} URLs.")
        except Exception as e:
            st.error(f"Sitemap load failed: {e}")

    if c2.button("Fetch sitemap from URL"):
        try:
            if sitemap_url.strip():
                urls = fetch_sitemap_urls_any(sitemap_url.strip())
                st.session_state["sitemap_urls"] = urls
                st.success(f"Loaded {len(urls)} URLs from URL.")
            else:
                st.warning("Provide a valid sitemap URL first.")
        except Exception as e:
            st.error(f"Sitemap fetch failed: {e}")

    # Linking knobs
    max_links = st.slider("Max auto-links / article", 0, 30, st.session_state.get("max_links", 8))
    st.session_state["max_links"] = max_links
    per_paragraph_max = st.slider("Max links / paragraph", 0, 10, st.session_state.get("per_paragraph_max", 2))
    st.session_state["per_paragraph_max"] = per_paragraph_max
    link_headings = st.checkbox("Also link inside headings", value=st.session_state.get("link_headings", False))
    st.session_state["link_headings"] = link_headings
    skip_lines_with_links = st.checkbox("Skip lines that already contain links", value=st.session_state.get("skip_lines_with_links", True))
    st.session_state["skip_lines_with_links"] = skip_lines_with_links
    hint_model_with_phrases = st.checkbox("Hint the model with internal phrases", value=st.session_state.get("hint_model_with_phrases", True))
    st.session_state["hint_model_with_phrases"] = hint_model_with_phrases

# Sidebar: WordPress Publish
with st.sidebar:
    st.subheader("📰 WordPress Publish")
    wp_site = st.text_input("Site URL", value=st.session_state.get("wp_site", "https://www.tastetorate.com"))
    st.session_state["wp_site"] = wp_site
    wp_user = st.text_input("WP Username (for App Password)", value=st.session_state.get("wp_user", ""))
    st.session_state["wp_user"] = wp_user
    wp_app_pw = st.text_input("WP Application Password", type="password", value=st.session_state.get("wp_app_pw", ""))
    st.session_state["wp_app_pw"] = wp_app_pw
    
    # Focus Keyword with Post status and Categories grouped together
    # wp_focus_keyword = st.text_input("Focus Keyword", value=st.session_state.get("wp_focus_keyword", ""))
    # st.session_state["wp_focus_keyword"] = wp_focus_keyword
    
    col1, col2 = st.columns(2)
    with col1:
        _status_options = ["draft", "publish"]
        _status_default = st.session_state.get("wp_status", "draft")
        wp_status = st.selectbox("Post status", _status_options, index=_status_options.index(_status_default) if _status_default in _status_options else 0)
        st.session_state["wp_status"] = wp_status
    with col2:
        wp_categories = st.text_input("Categories", value=st.session_state.get("wp_categories", "Recipes"))
        st.session_state["wp_categories"] = wp_categories
    
    wp_excerpt = st.text_area("Custom excerpt (optional)", height=80, value=st.session_state.get("wp_excerpt", ""))
    st.session_state["wp_excerpt"] = wp_excerpt
    
    # Test connection button
    if st.button("🔍 Test WordPress Connection"):
        site = st.session_state.get("wp_site", "").strip()
        user = st.session_state.get("wp_user", "").strip()
        app_pw = st.session_state.get("wp_app_pw", "").strip()
        
        if not (site and user and app_pw):
            st.warning("Please fill in Site URL, Username, and Application Password first.")
        else:
            with st.spinner("Testing connection..."):
                result = test_wp_connection(site, user, app_pw)
                if result["success"]:
                    st.success(result["message"])
                    if "api_info" in result:
                        st.caption(f"WordPress version: {result['api_info'].get('description', 'Unknown')}")
                else:
                    st.error(result["message"])
    
    st.caption("Tip: Use an Administrator Application Password so CPT endpoints/shortcodes work and HTML isn't stripped.")


# Sidebar: Save settings
with st.sidebar:
    if st.button("💾 Save Settings"):
        save_current_settings()

# ---------------------------------------------------------
# Main Panel
# ---------------------------------------------------------
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("✍️ Compose")
    focus_keyword = st.text_input("Focus Keyword", placeholder="e.g., black velvet cheesecake cookies")
    # Use focus_keyword as both topic and focus keyword for consistency
    topic = focus_keyword
    faq_text = st.text_area("Example FAQ (optional)", height=120, placeholder="Paste 4 example Q/As or leave blank")
    recipe_text = st.text_area("Full Recipe (parsed for Tasty + fallback)", height=220, placeholder="Paste recipe text (title, times, ingredients, instructions, notes, nutrition)", value=st.session_state.get("recipe_text", ""))
    st.session_state["recipe_text"] = recipe_text

    internal_phrase_hints = []
    if st.session_state.get("hint_model_with_phrases") and st.session_state.get("sitemap_urls"):
        try:
            idx = build_link_index(st.session_state["sitemap_urls"])
            internal_phrase_hints = [d["phrase"] for d in idx[:20]]
        except Exception:
            internal_phrase_hints = []

    if st.button("Generate Article"):
        if not topic.strip():
            st.warning("Please enter a topic before generating.")
        elif not focus_keyword.strip():
            st.warning("Please enter a focus keyword before generating.")
        else:
            # Create progress container
            progress_container = st.container()
            
            with progress_container:
                # Initialize progress bar and status
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                try:
                    model = st.session_state.get("model_name", "gpt-4.1")
                    temp = float(st.session_state.get("temperature", 0.6))
                    target = int(st.session_state.get("target_words", 1800))
                    
                    # Step 1: Generate SEO title
                    status_text.text("🎯 Generating SEO-optimized title...")
                    progress_bar.progress(10)
                    
                    seo_title = generate_seo_title(
                        topic=topic.strip(),
                        focus_keyword=focus_keyword.strip(),
                        model=model,
                        temperature=temp
                    )
                    
                    # Store SEO title in session state for WordPress publishing
                    st.session_state["seo_title"] = seo_title
                    
                    # Step 1.5: Generate comprehensive RankMath SEO metadata
                    status_text.text("🎯 Generating RankMath SEO metadata...")
                    progress_bar.progress(15)
                    
                    rankmath_metadata = generate_rankmath_seo_metadata(
                        topic=topic.strip(),
                        focus_keyword=focus_keyword.strip(),
                        article_content="",  # Will be updated after content generation
                        model=model
                    )
                    
                    # Store RankMath metadata in session state
                    st.session_state["rankmath_metadata"] = rankmath_metadata
                    
                    # Step 2: Generate article content
                    if st.session_state.get("use_multi_call", True):
                        status_text.text("✍️ Generating article content (multi-part)...")
                        progress_bar.progress(30)
                        
                        content_md = gpt_article_multi(
                            topic=topic.strip(),
                            faq_text=faq_text.strip(),
                            recipe_text=recipe_text.strip(),
                            focus_keyword=focus_keyword.strip(),
                            model=model,
                            phrase_hints=internal_phrase_hints if st.session_state.get("hint_model_with_phrases") else None,
                            target_words=target,
                            temperature=temp,
                            seo_title=seo_title,
                        )
                    else:
                        status_text.text("✍️ Generating article content (single call)...")
                        progress_bar.progress(30)
                        
                        content_md = gpt_article_single(
                            topic=topic.strip(),
                            faq_text=faq_text.strip(),
                            recipe_text=recipe_text.strip(),
                            focus_keyword=focus_keyword.strip(),
                            model=model,
                            phrase_hints=internal_phrase_hints if st.session_state.get("hint_model_with_phrases") else None,
                            temperature=temp,
                            seo_title=seo_title,
                        )

                    # Step 3: Add internal links
                    status_text.text("🔗 Adding internal links...")
                    progress_bar.progress(70)
                    
                    urls = st.session_state.get("sitemap_urls", [])
                    idx = build_link_index(filter_candidate_urls(urls)) if urls else []
                    content_linked = autolink_content(
                        content_md,
                        idx,
                        max_links=st.session_state.get("max_links", 14),
                        per_paragraph_max=st.session_state.get("per_paragraph_max", 3),
                        link_headings=st.session_state.get("link_headings", False),
                        skip_lines_with_links=st.session_state.get("skip_lines_with_links", True),
                        link_style=st.session_state.get("link_style", "html")
                    )

                    # Step 4: Add CTA
                    # status_text.text("📢 Adding call-to-action...")
                    # progress_bar.progress(90)
                    
                    # if st.session_state.get("append_cta"):
                    #     fb = st.session_state.get("fb_url") or ""
                    #     pin = st.session_state.get("pin_url") or ""
                    #     cta_bits = []
                    #     if fb: cta_bits.append(_wrap_link("Follow us on Facebook", _normalize_url(fb), st.session_state.get("link_style","html")))
                    #     if pin: cta_bits.append(_wrap_link("Find us on Pinterest", _normalize_url(pin), st.session_state.get("link_style","html")))
                    #     if cta_bits:
                    #         content_linked += "\n\n" + " · ".join(cta_bits)

                    # Step 5: Complete
                    status_text.text("✅ Article generation complete!")
                    progress_bar.progress(100)
                    
                    st.session_state["generated_md"] = content_linked
                    
                    # Update RankMath metadata with actual article content
                    if st.session_state.get("rankmath_metadata"):
                        updated_metadata = generate_rankmath_seo_metadata(
                            topic=topic.strip(),
                            focus_keyword=focus_keyword.strip(),
                            article_content=content_linked,
                            model=model
                        )
                        st.session_state["rankmath_metadata"] = updated_metadata
                    
                    # Clear progress indicators after a brief moment
                    import time
                    time.sleep(1)
                    progress_container.empty()
                    
                    # Show success message with article stats
                    word_count = len(content_linked.split())
                    char_count = len(content_linked)
                    st.success(f"🎉 Article generated successfully! ({word_count} words, {char_count} characters)")
                    
                except Exception as e:
                    progress_container.empty()
                    st.error(f"❌ Generation failed: {e}")

    if st.session_state.get("generated_md"):
        st.markdown("### Preview")
        
        # Get the article content
        content_to_display = st.session_state["generated_md"]
        
        # Add recipe card to preview if recipe text is provided
        recipe_text = st.session_state.get("recipe_text", "")
        if recipe_text and recipe_text.strip():
            try:
                parsed_recipe = parse_recipe_text_blocks(recipe_text)
                if parsed_recipe.get("title"):
                    # Add recipe card HTML to the content for preview
                    recipe_html = html_recipe_fallback(parsed_recipe)
                    content_to_display += "\n\n" + recipe_html
            except Exception:
                pass
        
        if st.session_state.get("link_style","html") == "html":
            # Convert Markdown to HTML for preview if links are HTML
            if recipe_text and recipe_text.strip():
                # For HTML preview with recipe card, use HTML component
                html_content = md_to_html(st.session_state["generated_md"])
                if "recipe_html" in locals():
                    html_content += "\n\n" + recipe_html
                st.components.v1.html(html_content, height=700, scrolling=True)
            else:
                st.components.v1.html(md_to_html(st.session_state["generated_md"]), height=700, scrolling=True)
        else:
            st.markdown(content_to_display)
        
        # RankMath SEO Metadata Display
        if st.session_state.get("rankmath_metadata"):
            st.markdown("---")
            st.markdown("### 🎯 RankMath SEO Metadata")
            st.markdown("Comprehensive SEO optimization data for your WordPress post.")
            
            metadata = st.session_state["rankmath_metadata"]
            
            # Create tabs for different metadata categories
            tab1, tab2, tab3, tab4 = st.tabs(["📝 Basic SEO", "🔗 Social Media", "📊 Schema Markup", "📋 Copy Text"])
            
            with tab1:
                st.markdown("#### Basic SEO Elements")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**SEO Title:**")
                    st.code(metadata.get('seo_title', 'N/A'), language=None)
                    
                    st.markdown("**Permalink/Slug:**")
                    st.code(metadata.get('permalink_slug', 'N/A'), language=None)
                    
                    st.markdown("**Focus Keyword:**")
                    st.code(metadata.get('focus_keyword', 'N/A'), language=None)
                
                with col2:
                    st.markdown("**Meta Description:**")
                    st.code(metadata.get('meta_description', 'N/A'), language=None)
                    
                    st.markdown("**Secondary Keywords:**")
                    keywords = metadata.get('secondary_keywords', [])
                    if keywords:
                        st.code(', '.join(keywords), language=None)
                    else:
                        st.code('N/A', language=None)
                    
                    st.markdown("**Additional Tags:**")
                    tags = metadata.get('additional_tags', [])
                    if tags:
                        st.code(', '.join(tags), language=None)
                    else:
                        st.code('N/A', language=None)
            
            with tab2:
                st.markdown("#### Social Media Optimization")
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Open Graph Title:**")
                    st.code(metadata.get('og_title', 'N/A'), language=None)
                    
                    st.markdown("**Open Graph Description:**")
                    st.code(metadata.get('og_description', 'N/A'), language=None)
                
                with col2:
                    st.markdown("**Twitter Title:**")
                    st.code(metadata.get('twitter_title', 'N/A'), language=None)
                    
                    st.markdown("**Twitter Description:**")
                    st.code(metadata.get('twitter_description', 'N/A'), language=None)
            
            with tab3:
                st.markdown("#### Schema Markup")
                schema_elements = metadata.get('schema_elements', {})
                
                st.markdown("**Schema Type:**")
                st.code(metadata.get('schema_type', 'Recipe'), language=None)
                
                if schema_elements:
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**Recipe Name:**")
                        st.code(schema_elements.get('name', 'N/A'), language=None)
                        
                        st.markdown("**Recipe Category:**")
                        st.code(schema_elements.get('recipeCategory', 'N/A'), language=None)
                    
                    with col2:
                        st.markdown("**Recipe Cuisine:**")
                        st.code(schema_elements.get('recipeCuisine', 'N/A'), language=None)
                        
                        st.markdown("**Keywords:**")
                        st.code(schema_elements.get('keywords', 'N/A'), language=None)
                    
                    st.markdown("**Schema Description:**")
                    st.code(schema_elements.get('description', 'N/A'), language=None)
            
            

# Generated: {metadata.get('generated_at', 'N/A')}"""
                
        st.info(f"📅 Generated: {metadata.get('generated_at', 'Unknown time')}")
        
        # Step 2: Image Prompt Generation
        st.markdown("---")
        st.markdown("### 🖼️ Step 2: Generate Image Prompts")
        st.markdown("Generate 7 MidJourney prompts with consistent styling and exact placement locations.")
        
        if st.button("🎨 Generate Image Prompts", key="generate_images"):
            if not st.session_state.get("generated_md"):
                st.warning("Generate the article first.")
            else:
                with st.spinner("Generating MidJourney prompts..."):
                    image_data = generate_midjourney_prompts(
                        article_content=st.session_state["generated_md"],
                        topic=topic or "Recipe",
                        focus_keyword=focus_keyword or "",
                        model=st.session_state.get("model_name", "gpt-4.1")
                    )
                    
                    if image_data and "images" in image_data:
                        st.session_state["image_prompts"] = image_data
                        st.success(f"✅ Generated {len(image_data['images'])} image prompts with seed: {image_data.get('seed', 'N/A')}")
                    else:
                        st.error("Failed to generate image prompts. Please try again.")
        
        # Display generated image prompts
        if st.session_state.get("image_prompts"):
            image_data = st.session_state["image_prompts"]
            
            st.markdown(f"**Batch Seed:** `{image_data.get('seed', 'N/A')}` (use this for all images)")
            st.markdown("**📋 Copy prompts in batches of 2 for MidJourney:**")
            
            # Group images for batch generation (2 at a time)
            images = image_data.get("images", [])
            batches = [
                ("Batch 1", images[0:2] if len(images) > 1 else images[0:1]),
                ("Batch 2", images[2:4] if len(images) > 3 else images[2:3] if len(images) > 2 else []),
                ("Batch 3", images[4:6] if len(images) > 5 else images[4:5] if len(images) > 4 else []),
                ("Batch 4", images[6:7] if len(images) > 6 else [])
            ]
            
            for batch_name, batch_images in batches:
                if batch_images:
                    with st.expander(f"🎯 {batch_name}: {' + '.join([img.get('type', '').title() for img in batch_images])}", expanded=True):
                        batch_prompts = []
                        for img in batch_images:
                            st.markdown(f"**{img.get('type', '').title()} Image**")
                            st.markdown(f"📍 *Placement:* {img.get('placement', 'N/A')}")
                            st.markdown(f"📝 *Description:* {img.get('description', 'N/A')}")
                            
                            prompt_text = img.get('prompt', '')
                            st.code(prompt_text, language="text")
                            batch_prompts.append(prompt_text)
                            st.markdown("---")
                        
                        # Copy all prompts in this batch
                        if batch_prompts:
                            combined_prompts = "\n\n".join(batch_prompts)
                            st.download_button(
                                 f"📋 Copy {batch_name} Prompts",
                                 data=combined_prompts,
                                 file_name=f"{to_slug(topic or 'recipe')}_{batch_name.lower().replace(' ', '_')}_prompts.txt",
                                 mime="text/plain",
                                 key=f"copy_{batch_name.lower().replace(' ', '_')}"
                             )
        
        # Step 3: Image Upload and Integration
        st.markdown("---")
        st.markdown("### 📤 Step 3: Upload Generated Images")
        st.markdown("Upload your generated MidJourney images and automatically insert them into the article at the specified locations.")
        
        if st.session_state.get("image_prompts"):
            images = st.session_state["image_prompts"].get("images", [])
            
            # Create file uploaders for each image type
            uploaded_images = {}
            
            st.markdown("**Upload images in the order they were generated:**")
            
            for i, img_info in enumerate(images):
                 img_type = img_info.get('type', f'image_{i+1}')
                 img_description = img_info.get('description', 'Image')
                 seo_metadata = img_info.get('seo_metadata', {})
                 
                 # Create expandable section for each image
                 with st.expander(f"📸 {img_type.title()} Image - Upload & SEO Metadata", expanded=False):
                     col_upload, col_seo = st.columns([1, 1])
                     
                     with col_upload:
                         st.markdown("**📤 Upload Image**")
                         uploaded_file = st.file_uploader(
                             f"{img_type.title()}",
                             type=["jpg", "jpeg", "png", "webp"],
                             key=f"upload_{img_type}_{i}",
                             help=f"Upload the {img_type} image"
                         )
                         
                         st.markdown(f"**📍 Placement:** {img_info.get('placement', 'N/A')}")
                         st.markdown(f"**📝 Description:** {img_description}")
                         
                         if uploaded_file:
                             uploaded_images[img_type] = {
                                 'file': uploaded_file,
                                 'placement': img_info.get('placement', ''),
                                 'description': img_description,
                                 'seo_metadata': seo_metadata
                             }
                     
                     with col_seo:
                         st.markdown("**🏷️ SEO Metadata**")
                         if seo_metadata:
                             st.markdown(f"**Alt Text:** {seo_metadata.get('alt_text', 'N/A')}")
                             st.markdown(f"**Filename:** {seo_metadata.get('filename', 'N/A')}")
                             st.markdown(f"**Caption:** {seo_metadata.get('caption', 'N/A')}")
                             st.markdown(f"**Description:** {seo_metadata.get('description', 'N/A')}")
                             
                             # Quality assurance checklist
                             st.markdown("**✅ Quality Checklist:**")
                             focus_keyword = st.session_state.get("image_prompts", {}).get("focus_keyword", "")
                             if focus_keyword:
                                 alt_has_keyword = focus_keyword.lower() in seo_metadata.get('alt_text', '').lower()
                                 filename_has_keyword = focus_keyword.lower().replace(' ', '-') in seo_metadata.get('filename', '').lower()
                                 
                                 st.markdown(f"{'✅' if alt_has_keyword else '❌'} Alt text includes keyword")
                                 st.markdown(f"{'✅' if filename_has_keyword else '❌'} Filename includes keyword")
                                 st.markdown(f"{'✅' if seo_metadata.get('caption') else '❌'} Caption provided")
                                 st.markdown(f"{'✅' if seo_metadata.get('description') else '❌'} Description provided")
                         else:
                             st.warning("No SEO metadata available. Please regenerate image prompts.")
            
            # Process and integrate images
            if uploaded_images and st.button("🔗 Insert Images into Article", key="insert_images"):
                try:
                    # Create a modified version of the article with image placeholders
                    modified_content = st.session_state["generated_md"]
                    
                    # Helper utilities for robust placement independent of upload order
                    import re
                    def _find_section_bounds(text: str, heading_variants):
                        for hv in heading_variants:
                            pos = text.find(hv)
                            if pos != -1:
                                start = pos + len(hv)
                                next_h = text.find("\n## ", start)
                                end = next_h if next_h != -1 else len(text)
                                return (pos, start, end, hv)
                        return (-1, -1, -1, None)

                    def _insert_at(text: str, idx: int, insert: str) -> str:
                        if idx < 0:
                            return text + insert
                        return text[:idx] + insert + text[idx:]

                    def _insert_after_nth_step(text: str, n: int, heading_variants):
                        # Find instructions section
                        s_pos, s_start, s_end, hv = _find_section_bounds(text, heading_variants)
                        if s_pos == -1:
                            return text  # no section, skip
                        section = text[s_start:s_end]
                        # Patterns for numbered lists (1. or 1) ) and explicit Step N labels
                        pat_num = re.compile(rf"^\s*{n}[\).]\s", re.MULTILINE)
                        pat_word = re.compile(rf"^\s*(?:Step\s+{n}\b)", re.IGNORECASE | re.MULTILINE)
                        m = pat_num.search(section) or pat_word.search(section)
                        if not m:
                            # Fallback placements: start/middle/end depending on n
                            if n == 1:
                                insert_idx = s_start
                            elif n == 2:
                                insert_idx = s_start + (s_end - s_start) // 2
                            else:
                                insert_idx = s_end
                            return _insert_at(text, insert_idx, img_content)
                        # determine the end of this step block -> until next step marker or blank line
                        step_start = s_start + m.end()
                        next_blank = re.search(r"\n\s*\n", text[step_start:s_end])
                        next_step = re.compile(r"^\s*(?:\d+[\).]\s|Step\s+\d+\b)", re.MULTILINE).search(text[step_start:s_end])
                        rel_end = None
                        if next_step:
                            rel_end = next_step.start()
                        if next_blank and (rel_end is None or next_blank.start() < rel_end):
                            rel_end = next_blank.end()
                        if rel_end is None:
                            insert_idx = s_end
                        else:
                            insert_idx = step_start + rel_end
                        return _insert_at(text, insert_idx, img_content)

                    # Deterministic expected order regardless of upload order
                    expected_order = ["ingredients", "step1", "step2", "step3", "serving", "recipe_card"]
                    ordered_types = [t for t in expected_order if t in uploaded_images]
                    for t in uploaded_images.keys():
                        if t not in ordered_types and t != "featured":
                            ordered_types.append(t)

                    # Insert images at appropriate locations based on the actual article structure
                    for img_type in ordered_types:
                        img_data = uploaded_images[img_type]
                        # Get SEO metadata
                        seo_metadata = img_data.get('seo_metadata', {})
                        file_name = seo_metadata.get('filename', img_data['file'].name)
                        alt_text = seo_metadata.get('alt_text', img_data['description'])
                        caption = seo_metadata.get('caption', img_data['description'])

                        # Create WordPress-compatible placeholder URLs (will be replaced with actual URLs during publishing)
                        current_date = datetime.datetime.now()
                        wp_placeholder_url = f"wp-content/uploads/{current_date.year}/{current_date.month:02d}/{file_name}"

                        # Use HTML for better WordPress compatibility
                        img_html = f'\n\n<img src="{wp_placeholder_url}" alt="{alt_text}" class="recipe-image" />\n<p class="image-caption"><em>{caption}</em></p>\n\n'
                        img_content = img_html

                        # Insert based on image type and the actual article structure
                        if img_type == "featured":
                            # Skip inserting featured image into article content - it will be set as WordPress featured image only
                            continue
                        elif img_type == "ingredients":
                            # Place immediately after Ingredients section ends
                            ing_headings = ["## Ingredients", "### Ingredients", "## Ingredient", "### Ingredient"]
                            pos, start, end, hv = _find_section_bounds(modified_content, ing_headings)
                            if pos != -1:
                                modified_content = _insert_at(modified_content, end, img_content)
                            else:
                                # fallback: before Instructions/How To Make It if present, else append
                                instr_heads = ["## Instructions", "## How To Make It", "## Method", "## Directions"]
                                i_pos, i_start, i_end, _ = _find_section_bounds(modified_content, instr_heads)
                                if i_pos != -1:
                                    modified_content = _insert_at(modified_content, i_pos, img_content)
                                else:
                                    modified_content += img_content
                        elif img_type == "step1":
                            instr_heads = ["## Instructions", "## How To Make It", "## Method", "## Directions"]
                            modified_content = _insert_after_nth_step(modified_content, 1, instr_heads)
                        elif img_type == "step2":
                            instr_heads = ["## Instructions", "## How To Make It", "## Method", "## Directions"]
                            modified_content = _insert_after_nth_step(modified_content, 2, instr_heads)
                        elif img_type == "step3":
                            instr_heads = ["## Instructions", "## How To Make It", "## Method", "## Directions"]
                            modified_content = _insert_after_nth_step(modified_content, 3, instr_heads)
                        elif img_type == "serving":
                            # Prefer Serving section, else before FAQ, else end
                            serving_heads = ["## Serving", "## Serving Suggestions", "## How to Serve"]
                            s_pos, s_start, s_end, _ = _find_section_bounds(modified_content, serving_heads)
                            if s_pos != -1:
                                modified_content = _insert_at(modified_content, s_start, img_content)
                            elif "## FAQ" in modified_content:
                                modified_content = modified_content.replace("## FAQ", img_content + "## FAQ")
                            else:
                                modified_content += img_content
                        elif img_type == "recipe_card" or "recipe" in img_type:
                            # Insert at the very end
                            modified_content += img_content
                        else:
                            # Unknown type -> append to end
                            modified_content += img_content
                    
                    # Update session state with modified content
                    st.session_state["generated_md_with_images"] = modified_content
                    
                    st.success(f"✅ Successfully inserted {len(uploaded_images)} images into the article!")
                    
                    # Show preview of modified content
                    with st.expander("📖 Preview Article with Images", expanded=True):
                        st.markdown(modified_content)
                    
                except Exception as e:
                    st.error(f"❌ Failed to insert images: {e}")
            
            # Download modified article
            if st.session_state.get("generated_md_with_images"):
                st.download_button(
                    "📄 Download Article with Images",
                    data=st.session_state["generated_md_with_images"],
                    file_name=f"{to_slug(topic or 'article')}_with_images.md",
                    mime="text/markdown",
                    key="download_with_images"
                )
                
                # Step 4: WordPress Publishing with Images
                st.markdown("---")
                st.markdown("### 🚀 Step 4: Publish to WordPress with Images")
                st.markdown("Upload all images with SEO metadata to WordPress and publish the complete article.")
                
                # Check WordPress credentials
                site = st.session_state.get("wp_site","").strip()
                user = st.session_state.get("wp_user","").strip()
                app_pw = st.session_state.get("wp_app_pw","").strip()
                status = st.session_state.get("wp_status","draft")
                
                if not (site and user and app_pw):
                    st.warning("⚠️ Please configure WordPress credentials in the sidebar first.")
                else:
                    # Show publishing options
                    col_pub1, col_pub2 = st.columns([2, 1])
                    
                    with col_pub1:
                        st.markdown(f"**📍 Target Site:** {site}")
                        st.markdown(f"**👤 User:** {user}")
                        st.markdown(f"**📊 Status:** {status.title()}")
                        
                        # Show images that will be uploaded
                        if uploaded_images:
                            st.markdown(f"**🖼️ Images to Upload:** {len(uploaded_images)}")
                            for img_type, img_data in uploaded_images.items():
                                seo_meta = img_data.get('seo_metadata', {})
                                st.markdown(f"• **{img_type.title()}**: {seo_meta.get('filename', img_data['file'].name)}")
                    
                    with col_pub2:
                        if st.button("🚀 Publish to WordPress", type="primary", use_container_width=True, key="publish_with_images"):
                            # Create progress container
                            wp_progress_container = st.container()
                            
                            with wp_progress_container:
                                wp_progress_bar = st.progress(0)
                                wp_status_text = st.empty()
                             
                                try:
                                    # Step 1: Upload all images with metadata
                                    wp_status_text.text("🖼️ Uploading images with SEO metadata...")
                                    wp_progress_bar.progress(10)
                                    
                                    uploaded_media = {}
                                    featured_media_id = None
                                    
                                    for i, (img_type, img_data) in enumerate(uploaded_images.items()):
                                        wp_status_text.text(f"📤 Uploading {img_type} image ({i+1}/{len(uploaded_images)})...")
                                        
                                        # Get file data
                                        file_bytes = img_data['file'].getvalue()
                                        seo_metadata = img_data.get('seo_metadata', {})
                                        
                                        # Use SEO filename if available, otherwise use original
                                        filename = seo_metadata.get('filename', img_data['file'].name)
                                        
                                        # Determine MIME type
                                        file_ext = filename.lower().split('.')[-1]
                                        mime_map = {
                                            'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                                            'png': 'image/png', 'webp': 'image/webp'
                                        }
                                        mime_type = mime_map.get(file_ext, 'image/jpeg')
                                        
                                        # Upload with metadata
                                        media_result = wp_upload_media_with_metadata(
                                            site, user, app_pw, file_bytes, filename, seo_metadata, mime_type
                                        )
                                        
                                        uploaded_media[img_type] = media_result
                                        
                                        # Set featured image
                                        if img_type == 'featured':
                                            featured_media_id = media_result['media_id']
                                        
                                        # Update progress
                                        progress = 10 + (i + 1) * 30 // len(uploaded_images)
                                        wp_progress_bar.progress(progress)
                                    
                                    wp_status_text.text(f"✅ All {len(uploaded_images)} images uploaded successfully!")
                                    wp_progress_bar.progress(50)
                                    
                                    # Step 2: Update article content with WordPress image URLs
                                    wp_status_text.text("🔗 Updating article with WordPress image URLs...")
                                    wp_progress_bar.progress(60)
                                    
                                    # Replace placeholder image URLs with actual WordPress URLs and enhance HTML
                                    final_content = st.session_state["generated_md_with_images"]
                                    for img_type, media_info in uploaded_media.items():
                                        # Get SEO metadata for this image
                                        img_data = uploaded_images[img_type]
                                        seo_metadata = img_data.get('seo_metadata', {})
                                        alt_text = seo_metadata.get('alt_text', img_data['description'])
                                        caption = seo_metadata.get('caption', img_data['description'])
                                        
                                        # Replace WordPress placeholder with actual WordPress URL
                                        current_date = datetime.datetime.now()
                                        placeholder = f"wp-content/uploads/{current_date.year}/{current_date.month:02d}/{media_info['filename']}"
                                        actual_url = media_info['url']
                                        
                                        # Create enhanced WordPress-optimized HTML
                                        wp_img_html = f'<img src="{actual_url}" alt="{alt_text}" class="recipe-image wp-image-{media_info["media_id"]}" style="max-width: 100%; height: auto; display: block; margin: 20px auto;" />'
                                        if caption:
                                            wp_img_html += f'<p class="image-caption" style="font-style: italic; text-align: center; margin: 10px 0; color: #666;">{caption}</p>'
                                        
                                        # Replace the entire HTML block with WordPress-optimized version
                                        old_html_pattern = f'<img src="{placeholder}" alt="{alt_text}" class="recipe-image" />\n<p class="image-caption"><em>{caption}</em></p>'
                                        final_content = final_content.replace(old_html_pattern, wp_img_html)
                                        
                                        # Also handle any remaining placeholder URLs
                                        final_content = final_content.replace(placeholder, actual_url)
                                    
                                    # Step 3: Process recipe if available
                                    wp_status_text.text("🍳 Processing recipe data...")
                                    wp_progress_bar.progress(70)
                                    
                                    tasty_id = None
                                    try:
                                        parsed_recipe = parse_recipe_text_blocks(recipe_text or "")
                                        if parsed_recipe.get("title"):
                                            tasty_id = create_tasty_recipe_via_rest(site, user, app_pw, parsed_recipe)
                                    except Exception:
                                        tasty_id = None
                                    
                                    # Step 4: Convert to HTML and create post
                                    wp_status_text.text("📝 Converting content and creating post...")
                                    wp_progress_bar.progress(80)
                                    
                                    # Convert markdown to HTML
                                    content_html = md_to_html(final_content)
                                    
                                    # Add recipe shortcode if available
                                    if tasty_id:
                                        recipe_shortcode = embed_tasty_recipe_shortcode(tasty_id)
                                        content_html += f"\n\n{recipe_shortcode}"
                                    
                                    # Handle categories and tags
                                    cat_ids = None
                                    try:
                                        if st.session_state.get("wp_categories"):
                                            cat_names = [c.strip() for c in st.session_state["wp_categories"].split(",") if c.strip()]
                                            cat_ids = ensure_terms(site, user, app_pw, "categories", cat_names)
                                    except Exception:
                                        pass
                                    
                                    # Create the post
                                    wp_status_text.text("🚀 Publishing to WordPress...")
                                    wp_progress_bar.progress(90)
                                    
                                    seo_title = st.session_state.get("seo_title")
                                    wp_title = seo_title if seo_title else (topic or "Untitled")
                                    # Auto-generate slug from focus keyword
                                    focus_keyword = st.session_state.get("focus_keyword", "")
                                    slug_val = to_slug(focus_keyword or topic or "") or None
                                    excerpt_val = st.session_state.get("wp_excerpt") or excerpt_from_text(content_html, 40)
                                    
                                    post = wp_create_post(
                                        site, user, app_pw,
                                        title=wp_title,
                                        content_html=content_html,
                                        status=status,
                                        category_ids=cat_ids,
                                        tag_ids=None,
                                        featured_media_id=featured_media_id,
                                        excerpt=excerpt_val,
                                        slug=slug_val
                                    )
                                    
                                    # Success!
                                    wp_status_text.text("✅ Article published successfully with all images!")
                                    wp_progress_bar.progress(100)
                                    
                                    # Clear progress after brief moment
                                    import time
                                    time.sleep(2)
                                    wp_progress_container.empty()
                                    
                                    # Show success message with details
                                    st.success(f"🎉 **Article Published Successfully!**\n\n"
                                             f"📄 **Post:** [{wp_title}]({post.get('link', '#')})\n\n"
                                             f"🖼️ **Images Uploaded:** {len(uploaded_images)}\n\n"
                                             f"📊 **Status:** {status.title()}")
                                    
                                    # Persist publish state for Step 4 unlock
                                    st.session_state["wp_post_published"] = True
                                    st.session_state["wp_last_post"] = post
                                    st.session_state["wp_last_post_link"] = post.get('link', '#')
                                    
                                    # Show uploaded images summary
                                    with st.expander("📋 Uploaded Images Summary", expanded=False):
                                        for img_type, media_info in uploaded_media.items():
                                            st.markdown(f"**{img_type.title()}:**")
                                            st.markdown(f"• URL: {media_info['url']}")
                                            st.markdown(f"• Filename: {media_info['filename']}")
                                            seo_meta = media_info['seo_metadata']
                                            if seo_meta.get('alt_text'):
                                                st.markdown(f"• Alt Text: {seo_meta['alt_text']}")
                                            st.markdown("---")
                                    
                                except Exception as e:
                                    wp_progress_container.empty()
                                    st.error(f"❌ Publishing failed: {e}")
        else:
            st.info("Generate image prompts first to enable image upload functionality.")

with col2:
    st.subheader("📋 Export & Publish Workflow")
    
    # Check if article is generated
    article_ready = bool(st.session_state.get("generated_md"))
    
    # Step 1: Article Generation Status
    with st.container():
        if article_ready:
            st.success("✅ **Step 1:** Article Generated")
        else:
            st.info("📝 **Step 1:** Generate Article First")
            st.caption("Complete the article generation before proceeding")
    
    st.divider()
    
    # Step 2: Export Options
    with st.container():
        st.markdown("📄 **Step 2:** Export Options")
        
        if article_ready:
            col_export1, col_export2 = st.columns(2)
            
            with col_export1:
                if st.button("📥 Download DOCX", use_container_width=True):
                    bio = build_docx_from_content(topic or "Article", st.session_state["generated_md"])
                    st.download_button(
                        "💾 Save DOCX File",
                        data=bio,
                        file_name=f"{to_slug(topic or 'article')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with col_export2:
                if st.button("📋 Copy Markdown", use_container_width=True):
                    st.code(st.session_state["generated_md"], language="markdown")
        else:
            st.caption("⏳ Export options available after article generation")
    
    st.divider()
    
    # Step 3: WordPress Publishing
    with st.container():
        st.markdown("🚀 **Step 3:** WordPress Publishing")
        
        # WordPress connection status
        wp_configured = bool(
            st.session_state.get("wp_site", "").strip() and 
            st.session_state.get("wp_user", "").strip() and 
            st.session_state.get("wp_app_pw", "").strip()
        )
        
        if not wp_configured:
            st.warning("⚙️ Configure WordPress settings in sidebar first")
        elif not article_ready:
            st.info("📝 Generate article before publishing")
        else:
            st.success("✅ Ready to publish to WordPress")
            
            if st.button("🚀 Publish to WordPress", type="primary", use_container_width=True):
                if not st.session_state.get("generated_md"):
                    st.warning("Generate the article first.")
                else:
                    site = st.session_state.get("wp_site","").strip()
                    user = st.session_state.get("wp_user","").strip()
                    app_pw = st.session_state.get("wp_app_pw","").strip()
                    status = st.session_state.get("wp_status","draft")
                    
                    if not (site and user and app_pw):
                        st.warning("Please fill WP site, username, and application password.")
                    else:
                        # Create progress container for WordPress posting
                         wp_progress_container = st.container()
                         
                         with wp_progress_container:
                             wp_progress_bar = st.progress(0)
                             wp_status_text = st.empty()
                    
                    try:
                        # Step 1: Upload featured image (if provided)
                        wp_status_text.text("🖼️ Processing featured image...")
                        wp_progress_bar.progress(20)
                        
                        feat_media_id = None
                        try:
                            feat_file_from_session = st.session_state.get("feat_file")
                            if feat_file_from_session is not None:
                                file_bytes = feat_file_from_session.getvalue()
                                feat_media_id = wp_upload_media(site, user, app_pw, file_bytes, feat_file_from_session.name)
                                wp_status_text.text("✅ Featured image uploaded successfully")
                            elif st.session_state.get("feat_url"):
                                rimg = requests.get(st.session_state["feat_url"], timeout=20)
                                rimg.raise_for_status()
                                fname = st.session_state["feat_url"].split("/")[-1] or "featured.jpg"
                                feat_media_id = wp_upload_media(site, user, app_pw, rimg.content, fname)
                                wp_status_text.text("✅ Featured image downloaded and uploaded")
                        except Exception as e:
                            wp_status_text.text(f"⚠️ Featured image skipped: {e}")
                            import time
                            time.sleep(1)

                        # Step 2: Process recipe
                        wp_status_text.text("🍳 Processing recipe data...")
                        wp_progress_bar.progress(40)
                        
                        tasty_id = None
                        try:
                            parsed_recipe = None
                            rtxt = recipe_text or ""
                            if rtxt.strip():
                                parsed_recipe = parse_recipe_text_blocks(rtxt)
                            if (not parsed_recipe or not (parsed_recipe.get("ingredients") or parsed_recipe.get("instructions"))):
                                md = st.session_state.get("generated_md", "")
                                if md.strip():
                                    extracted = _extract_recipe_from_article_md(md)
                                    if extracted and (extracted.get("ingredients") or extracted.get("instructions")):
                                        parsed_recipe = extracted
                            if not parsed_recipe or not (parsed_recipe.get("ingredients") or parsed_recipe.get("instructions")):
                                parsed_recipe = synthesize_recipe_from_context(
                                    topic=st.session_state.get("topic",""),
                                    focus_keyword=st.session_state.get("focus_keyword",""),
                                    full_recipe_text=rtxt,
                                    article_md=st.session_state.get("generated_md","")
                                )
                            if parsed_recipe and (parsed_recipe.get("ingredients") or parsed_recipe.get("instructions")):
                                parsed_recipe = _make_recipe_halal(parsed_recipe)
                                tasty_id = create_tasty_recipe_via_rest(site, user, app_pw, parsed_recipe)
                        except Exception:
                            tasty_id = None

                        # Step 3: Convert content to HTML
                        wp_status_text.text("📝 Converting content to HTML...")
                        wp_progress_bar.progress(60)
                        
                        content_md = st.session_state["generated_md"]
                        content_html = md_to_html(content_md)
                        
                        if tasty_id:
                            content_html += "\n\n" + embed_tasty_recipe_shortcode(tasty_id)
                        else:
                            try:
                                rtxt = recipe_text or ""
                                parsed = None
                                if rtxt.strip():
                                    parsed = parse_recipe_text_blocks(rtxt)
                                if (not parsed or not (parsed.get("ingredients") or parsed.get("instructions"))):
                                    md = st.session_state.get("generated_md", "")
                                    if md.strip():
                                        extracted = _extract_recipe_from_article_md(md)
                                        if extracted and (extracted.get("ingredients") or extracted.get("instructions")):
                                            parsed = extracted
                                if not parsed or not (parsed.get("ingredients") or parsed.get("instructions")):
                                    parsed = synthesize_recipe_from_context(
                                        topic=st.session_state.get("topic",""),
                                        focus_keyword=st.session_state.get("focus_keyword",""),
                                        full_recipe_text=rtxt,
                                        article_md=st.session_state.get("generated_md","")
                                    )
                                if parsed and (parsed.get("ingredients") or parsed.get("instructions")):
                                    parsed = _make_recipe_halal(parsed)
                                    content_html += "\n\n" + html_recipe_fallback(parsed)
                            except Exception:
                                pass

                        # Step 4: Setup categories and tags
                        wp_status_text.text("🏷️ Setting up categories and tags...")
                        wp_progress_bar.progress(80)
                        
                        cat_ids = ensure_terms(site, user, app_pw, "categories", [c.strip() for c in st.session_state.get("wp_categories","").split(",") if c.strip()])
                        
                        # Auto-generate slug from focus keyword
                        focus_keyword = st.session_state.get("focus_keyword", "")
                        slug_val = to_slug(focus_keyword or topic or "") or None
                        excerpt_val = st.session_state.get("wp_excerpt") or excerpt_from_text(content_html, 40)

                        # Step 5: Create WordPress post
                        wp_status_text.text("📤 Publishing to WordPress...")
                        wp_progress_bar.progress(95)
                        
                        # Use the generated SEO title for WordPress post
                        seo_title = st.session_state.get("seo_title")
                        wp_title = seo_title if seo_title else (topic or "Untitled")
                        
                        post = wp_create_post(
                            site, user, app_pw,
                            title=wp_title,
                            content_html=content_html,
                            status=status,
                            category_ids=cat_ids or None,
                            tag_ids=None,
                            featured_media_id=feat_media_id,
                            excerpt=excerpt_val,
                            slug=slug_val
                        )
                        
                        # Step 6: Complete
                        wp_status_text.text("✅ WordPress post published successfully!")
                        wp_progress_bar.progress(100)
                        
                        # Clear progress indicators after a brief moment
                        import time
                        time.sleep(1)
                        wp_progress_container.empty()
                        
                        # Persist publish state for Step 4 unlock
                        st.session_state["wp_post_published"] = True
                        st.session_state["wp_last_post"] = post
                        st.session_state["wp_last_post_link"] = post.get('link', '#')
                        
                        st.success(f"🎉 WordPress post created successfully! [View Post]({post.get('link','#')})")
                        
                    except Exception as e:
                        wp_progress_container.empty()
                        st.error(f"❌ Publishing failed: {e}")

# Step 4: JavaScript Tasty Recipe Card (unlocked after publishing)
with col2:
    st.divider()
    with st.container():
        st.markdown("🍫 **Step 4:** TASTY RECIPE CARD (Pure JavaScript Version)")
        if not st.session_state.get("wp_post_published"):
            st.info("Publish your article in Step 3 to unlock this step.")
        else:
            st.caption("Full Recipe is optional — we'll use it if provided, otherwise we'll extract from your article or synthesize a well-organized recipe card.")
            if st.button("Generate TASTY RECIPE CARD", use_container_width=True):
                rtxt = st.session_state.get("recipe_text", "")
                try:
                    recipe = None
                    if rtxt.strip():
                        recipe = parse_recipe_text_blocks(rtxt)
                    if (not recipe or not (recipe.get("ingredients") or recipe.get("instructions"))):
                        md = st.session_state.get("generated_md", "")
                        if md.strip():
                            extracted = _extract_recipe_from_article_md(md)
                            if extracted and (extracted.get("ingredients") or extracted.get("instructions")):
                                recipe = extracted
                    if not recipe or not (recipe.get("ingredients") or recipe.get("instructions")):
                        recipe = synthesize_recipe_from_context(
                            topic=st.session_state.get("topic",""),
                            focus_keyword=st.session_state.get("focus_keyword",""),
                            full_recipe_text=rtxt,
                            article_md=st.session_state.get("generated_md","")
                        )
                    if recipe and (recipe.get("ingredients") or recipe.get("instructions")):
                        recipe = _make_recipe_halal(recipe)
                        normalized = _normalize_recipe_for_tasty(recipe, author_name=st.session_state.get("author_name", ""))
                        st.session_state["js_recipe_card"] = generate_js_recipe_card(normalized)
                        st.success("JavaScript snippet generated below. Copy and paste it into your browser console.")
                    else:
                        st.warning("No recipe detected from your inputs. Paste a recipe in 'Full Recipe' or ensure your article has a clear recipe section.")
                except Exception as e:
                    st.error(f"Could not generate JS snippet: {e}")
            if st.session_state.get("js_recipe_card"):
                st.text_area("Copy & Paste the JavaScript below", value=st.session_state["js_recipe_card"], height=260, key="js_recipe_card_full_textarea")
                st.download_button("💾 Download JS file", data=st.session_state["js_recipe_card"], file_name="tasty_recipe_fill.js", mime="text/javascript", use_container_width=True, key="download_js_full")
                components.html(f"""
<div style=\"display:flex; justify-content:flex-end; margin-top:6px;\">
  <button onclick=\"navigator.clipboard.writeText(document.getElementById('js_full_copy').value); this.innerText='Copied!'; setTimeout(()=>this.innerText='📋 Copy JavaScript',1500);\" style=\"padding:6px 10px; font-size:14px;\">📋 Copy JavaScript</button>
</div>
<textarea id=\"js_full_copy\" style=\"position:absolute; left:-9999px; top:-9999px;\">{html.escape(st.session_state["js_recipe_card"])}</textarea>
""", height=40)
                st.caption("Tip: Paste into your browser console on the Tasty Recipes edit page, then press Enter. This fills all fields including description, ingredients, instructions, notes, times, yield, tags, and nutrition.")

# Gentle nudge for auto-post (explicit click keeps control clear)
def _normalize_recipe_for_tasty(recipe: dict, author_name: str = None) -> dict:
    try:
        out = dict(recipe or {})
    except Exception:
        out = {}
    # Title and author
    out["title"] = (
        out.get("title")
        or st.session_state.get("seo_title")
        or st.session_state.get("focus_keyword")
        or "Untitled Recipe"
    )
    if author_name:
        out["author"] = author_name
    else:
        out["author"] = out.get("author") or st.session_state.get("author_name") or ""

    # Description
    out["description"] = out.get("description") or ""

    # Normalize list-like fields
    def ensure_list(x):
        if x is None:
            return []
        if isinstance(x, (list, tuple)):
            return [str(i).strip() for i in x if str(i).strip()]
        if isinstance(x, str):
            return [s.strip("•- \t") for s in x.splitlines() if s.strip()]
        try:
            return [str(x).strip()]
        except Exception:
            return []

    out["ingredients"] = ensure_list(out.get("ingredients"))
    out["instructions"] = ensure_list(out.get("instructions"))

    notes = out.get("notes")
    if isinstance(notes, (list, tuple)):
        out["notes"] = "\n".join(str(n).strip() for n in notes if str(n).strip())
    else:
        out["notes"] = notes or ""

    # Details fields with safe defaults
    for key in [
        "prep_time",
        "cook_time",
        "total_time",
        "yield",
        "category",
        "method",
        "cuisine",
        "diet",
        "keywords",
        "serving_size",
    ]:
        if out.get(key) is None:
            out[key] = ""

    # Normalize keywords to comma string
    if isinstance(out.get("keywords"), (list, tuple)):
        out["keywords"] = ", ".join([str(k).strip() for k in out["keywords"] if str(k).strip()])

    # Normalize nutrition with aliases, preserve originals, add display-friendly keys, and mirror to top-level
    nut = out.get("nutrition") or {}
    aliases = {
        "serving_size": ["serving_size", "Serving Size", "Serving", "ServingSize", "Serving size"],
        "calories": ["calories", "Calories", "kcal", "Kcal", "KCAL"],
        "sugar": ["sugar", "sugars", "Sugar", "Sugars"],
        "sodium": ["sodium", "salt", "Sodium", "Salt"],
        "fat": ["fat", "total_fat", "Total Fat", "Fat", "total fat"],
        "saturated_fat": ["saturated_fat", "saturated", "Saturated Fat"],
        "unsaturated_fat": [
            "unsaturated_fat",
            "unsaturated",
            "Unsaturated Fat",
            "polyunsaturated_fat",
            "monounsaturated_fat",
            "Polyunsaturated Fat",
            "Monounsaturated Fat",
        ],
        "trans_fat": ["trans_fat", "trans", "Trans Fat"],
        "carbohydrates": [
            "carbohydrates",
            "carbs",
            "total_carbohydrates",
            "Total Carbohydrates",
            "Carbohydrates",
            "Total Carbs",
            "total carbs",
        ],
        "fiber": ["fiber", "dietary_fiber", "Dietary Fiber", "Fiber"],
        "protein": ["protein", "proteins", "Protein"],
        "cholesterol": ["cholesterol", "Cholesterol"],
    }

    # Preserve original nutrition keys
    preserved = {}
    try:
        preserved.update({str(k): v for k, v in nut.items()})
    except Exception:
        preserved = dict(nut)

    norm_nut = {}
    for canon, alts in aliases.items():
        val = out.get(canon)
        if not val:
            val = nut.get(canon)
        if not val:
            for a in alts:
                if out.get(a):
                    val = out.get(a)
                    break
                if nut.get(a):
                    val = nut.get(a)
                    break
        norm_nut[canon] = str(val).strip() if (val is not None and str(val).strip()) else ""

    # Merge preserved + canonical + display-friendly
    preserved.update(norm_nut)
    display_keys = {
        "serving_size": "Serving Size",
        "calories": "Calories",
        "sugar": "Sugar",
        "sodium": "Sodium",
        "fat": "Total Fat",
        "saturated_fat": "Saturated Fat",
        "unsaturated_fat": "Unsaturated Fat",
        "trans_fat": "Trans Fat",
        "carbohydrates": "Total Carbohydrates",
        "fiber": "Dietary Fiber",
        "protein": "Protein",
        "cholesterol": "Cholesterol",
    }
    for canon, pretty in display_keys.items():
        if norm_nut.get(canon):
            preserved.setdefault(pretty, norm_nut[canon])

    out["nutrition"] = preserved
    # Mirror to top-level for convenience
    out.update(norm_nut)
    # Set top-level serving_size from nutrition if missing
    if not out.get("serving_size") and norm_nut.get("serving_size"):
        out["serving_size"] = norm_nut["serving_size"]

    return out


st.divider()
with st.container():
    st.markdown("---")
    st.markdown("### 🚀 Step 5: AI-Generated TASTY RECIPE CARD (Pure JavaScript)")
    st.caption("This step sends your article and optional 'Full Recipe (parsed for Tasty + fallback)' to GPT to synthesize a well-organized recipe tailored for the WP Tasty plugin. It does not copy your raw text verbatim; it restructures and cleans it specifically for the card.")
    if st.button("Generate via GPT (Best for Tasty)", use_container_width=True, key="generate_tasty_js_ai"):
        try:
            rtxt = st.session_state.get("recipe_text", "")
            md = st.session_state.get("generated_md", "")
            recipe = synthesize_recipe_from_context(
                topic=st.session_state.get("topic", ""),
                focus_keyword=st.session_state.get("focus_keyword", ""),
                full_recipe_text=rtxt,
                article_md=md
            )
            if recipe and (recipe.get("ingredients") or recipe.get("instructions")):
                recipe = _make_recipe_halal(recipe)
                normalized = _normalize_recipe_for_tasty(recipe, author_name=st.session_state.get("author_name", ""))
                st.session_state["js_recipe_card_ai"] = generate_js_recipe_card(normalized)
                st.success("AI-generated JavaScript snippet ready below. Paste it into your browser console on the Tasty Recipes edit page.")
            else:
                st.warning("AI could not synthesize a recipe from your inputs. Please ensure your article includes recipe details or provide them in 'Full Recipe'.")
        except Exception as e:
            st.error(f"Could not generate AI JS snippet: {e}")

    if st.session_state.get("js_recipe_card_ai"):
        st.text_area("Copy & Paste the JavaScript below (AI)", value=st.session_state["js_recipe_card_ai"], height=260, key="js_recipe_card_ai_textarea")
        st.download_button("💾 Download AI JS file", data=st.session_state["js_recipe_card_ai"], file_name="tasty_recipe_fill_ai.js", mime="text/javascript", use_container_width=True, key="download_js_ai")
#         components.html(f"""
# <div style=\"display:flex; justify-content:flex-end; margin-top:6px;\">
#   <button onclick=\"navigator.clipboard.writeText(document.getElementById('js_ai_copy').value); this.innerText='Copied!'; setTimeout(()=>this.innerText='📋 Copy JavaScript (AI)',1500);\" style=\"padding:6px 10px; font-size:14px;\">📋 Copy JavaScript (AI)</button>
# </div>
# <textarea id=\"js_ai_copy\" style=\"position:absolute; left:-9999px; top:-9999px;\">{html.escape(st.session_state["js_recipe_card_ai"])}</textarea>
# """, height=40)
        st.caption("Tip: Paste into your browser console on the Tasty Recipes edit page, then press Enter. This fills description, ingredients, instructions, notes, times, yield, tags, and nutrition.")

# --- Step 6: Generate Recipe JSON-LD (Schema.org) from the published post ---
st.markdown("---")
st.markdown("### 🚀 Step 6: Recipe JSON-LD (Schema.org)")
st.caption("Only includes fields we can confidently extract. No fake data. Paste your published post URL below.")

# Helper: find Recipe JSON-LD already present in the page (if any)
def _find_recipe_jsonld_in_html(html_text: str):
    try:
        scripts = re.findall(r"<script[^>]+type=\"(?:application|text)/ld\+json\"[^>]*>(.*?)</script>", html_text, flags=re.I | re.S)
        for block in scripts:
            block = block.strip()
            if not block:
                continue
            try:
                data = json.loads(block)
            except Exception:
                # Some pages include multiple JSON objects concatenated; attempt a naive fix by wrapping in [] and splitting
                parts = re.split(r"}\s*,\s*{", block)
                if len(parts) > 1:
                    try:
                        data = json.loads(f"[{block}]")
                    except Exception:
                        continue
                else:
                    continue
            # Normalize to list for iteration
            candidates = data if isinstance(data, list) else [data]
            for item in candidates:
                if isinstance(item, dict):
                    atype = item.get("@type")
                    if isinstance(atype, list):
                        if any(str(t).lower() == "recipe" for t in atype):
                            return item
                    elif isinstance(atype, str) and atype.lower() == "recipe":
                        return item
    except Exception:
        pass
    return None

# Helper: extract OpenGraph image from page
def _extract_og_image(html_text: str) -> str:
    m = re.search(r"<meta[^>]+property=\"og:image\"[^>]+content=\"([^\"]+)\"", html_text, flags=re.I)
    if m:
        return m.group(1)
    return ""

# Helper: normalize lists for schema
def _norm_list(val):
    if not val:
        return []
    if isinstance(val, list):
        return [str(x).strip() for x in val if str(x).strip()]
    if isinstance(val, str):
        return [x.strip() for x in re.split(r"\n+|;", val) if x.strip()]
    return []

# Helper: Build Recipe JSON-LD using only known data
def _build_recipe_schema(recipe: dict, page_url: str = "", name: str = "", description: str = "", image: str = "", author: str = "", keywords: str = "") -> dict:
    schema = {"@context": "https://schema.org", "@type": "Recipe"}

    if name:
        schema["name"] = name
    if description:
        schema["description"] = description
    if image:
        schema["image"] = image
    if author:
        schema["author"] = {"@type": "Person", "name": author}
    if keywords:
        schema["keywords"] = keywords

    # Yield / servings
    ry = recipe.get("yield") or recipe.get("servings") or recipe.get("recipe_yield")
    if ry:
        schema["recipeYield"] = str(ry)

    # Times (minutes to ISO 8601)
    def _mins(*keys):
        for k in keys:
            v = recipe.get(k)
            if isinstance(v, (int, float)):
                return int(v)
            if isinstance(v, str):
                # extract digits
                m = re.search(r"(\d+)", v)
                if m:
                    return int(m.group(1))
        return None

    pt = _mins("prep_time_min", "prepTimeMinutes", "prep_time")
    ct = _mins("cook_time_min", "cookTimeMinutes", "cook_time")
    tt = _mins("total_time_min", "totalTimeMinutes", "total_time")
    from_minutes = _to_iso8601_minutes  # already defined above in this file
    if pt is not None:
        iso = from_minutes(pt)
        if iso:
            schema["prepTime"] = iso
    if ct is not None:
        iso = from_minutes(ct)
        if iso:
            schema["cookTime"] = iso
    if tt is not None:
        iso = from_minutes(tt)
        if iso:
            schema["totalTime"] = iso

    # Ingredients / Instructions
    ings = _norm_list(recipe.get("ingredients") or recipe.get("recipeIngredient"))
    if ings:
        schema["recipeIngredient"] = ings

    instr = _norm_list(recipe.get("instructions") or recipe.get("recipeInstructions"))
    if instr:
        # If already objects, keep them; else convert to HowToStep
        if instr and isinstance(instr[0], dict):
            schema["recipeInstructions"] = instr
        else:
            schema["recipeInstructions"] = [{"@type": "HowToStep", "text": s} for s in instr]

    # Nutrition
    nut = recipe.get("nutrition")
    if isinstance(nut, dict):
        # Keep only known simple fields
        keep = {k: v for k, v in nut.items() if v}
        if keep:
            schema["nutrition"] = keep

    # mainEntityOfPage
    if page_url:
        schema["mainEntityOfPage"] = {"@type": "WebPage", "@id": page_url}

    return schema

# UI: Only enable when we have a published link
post_url_default = st.session_state.get("wp_last_post_link", "")
post_url = st.text_input("Published Post URL", value=post_url_default, key="schema_post_url", placeholder="https://example.com/your-post/")

if st.button("Extract & Generate Schema JSON-LD", key="build_recipe_schema_btn", use_container_width=True):
    if not post_url.strip():
        st.warning("Please paste the published post URL first.")
    else:
        try:
            with st.spinner("Fetching page and extracting data..."):
                page_html = ""
                og_image = ""
                try:
                    r = requests.get(post_url.strip(), timeout=20)
                    r.raise_for_status()
                    page_html = r.text
                    og_image = _extract_og_image(page_html)
                except Exception:
                    page_html = ""
                # 1) Try to use existing Recipe JSON-LD from the page (no fake data)
                page_schema = _find_recipe_jsonld_in_html(page_html) if page_html else None

                if page_schema:
                    # Keep as-is but ensure essential context/type exist
                    schema_obj = {"@context": "https://schema.org", "@type": "Recipe"}
                    # Merge with page_schema fields
                    for k, v in page_schema.items():
                        if v is not None and k not in ("@context",):
                            schema_obj[k] = v
                    # Ensure image present if missing and we have og:image
                    if "image" not in schema_obj and og_image:
                        schema_obj["image"] = og_image
                else:
                    # 2) Derive from our article and any recipe we can parse (no hallucination)
                    recipe = None
                    rtxt = st.session_state.get("recipe_text", "")
                    if rtxt.strip():
                        try:
                            recipe = parse_recipe_text_blocks(rtxt)
                        except Exception:
                            recipe = None
                    if (not recipe) and st.session_state.get("generated_md", "").strip():
                        try:
                            extracted = _extract_recipe_from_article_md(st.session_state["generated_md"])  # may be empty
                            if extracted and (extracted.get("ingredients") or extracted.get("instructions")):
                                recipe = extracted
                        except Exception:
                            recipe = None
                    if not recipe:
                        recipe = {}

                    # Compose context values strictly from available data
                    name = st.session_state.get("seo_title") or st.session_state.get("topic", "")
                    desc = st.session_state.get("wp_excerpt")
                    if not desc:
                        try:
                            content_md = st.session_state.get("generated_md", "")
                            if content_md:
                                desc = excerpt_from_text(md_to_html(content_md), 40)
                        except Exception:
                            desc = ""
                    image_url = og_image or st.session_state.get("feat_url", "")
                    author_name = st.session_state.get("author_name", "")
                    kw = st.session_state.get("focus_keyword", "")

                    schema_obj = _build_recipe_schema(
                        recipe=recipe,
                        page_url=post_url.strip(),
                        name=name or "",
                        description=desc or "",
                        image=image_url or "",
                        author=author_name or "",
                        keywords=kw or "",
                    )

                # Store and show
                schema_json = json.dumps(schema_obj, ensure_ascii=False, indent=2)
                st.session_state["recipe_schema_jsonld"] = schema_json

                # Save to specified markdown file path
                out_path = "/Users/useraccount/Documents/Blogging/py agents/shemamarkup.md"
                try:
                    with open(out_path, "w", encoding="utf-8") as f:
                        f.write(schema_json)
                    st.success(f"Saved recipe schema to: {out_path}")
                except Exception as e:
                    st.warning(f"Could not write to {out_path}: {e}")

                st.success("Recipe JSON-LD generated.")
        except Exception as e:
            st.error(f"Failed to generate schema: {e}")

if st.session_state.get("recipe_schema_jsonld"):
    st.text_area("Generated JSON-LD", value=st.session_state["recipe_schema_jsonld"], height=260, key="schema_jsonld_textarea")
    st.download_button(
        "💾 Download recipe.schema.json",
        data=st.session_state["recipe_schema_jsonld"].encode("utf-8"),
        file_name="recipe.schema.json",
        mime="application/ld+json",
        use_container_width=True,
        key="download_recipe_schema_jsonld",
    )

if st.session_state.get("auto_post") and st.session_state.get("generated_md"):
    st.info("Auto-post is enabled. Click 'Post to WordPress' to publish the current draft.")